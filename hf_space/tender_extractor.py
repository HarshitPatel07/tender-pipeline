"""
tender_extractor.py
===================
Free, zero-cost tender document summariser for ADCA India.

Walks a Google Drive folder of tender sub-folders, reads every tender-issued
document (PDF / DOCX / XLS, OCR-ing scanned pages), and extracts 13 summary
fields per tender into an Excel workbook with page-level evidence.

Two extraction layers:
  1. RULES  - deterministic label/regex + section capture. Offline, free,
              near-perfect on the fixed GeM bid template.
  2. GEMINI - free-tier LLM pass for the judgement fields (Purpose, Eligibility,
              Scope, Penalty) and to fill whatever rules missed.

Where the two layers disagree on a money/date field, BOTH are reported and the
cell is flagged CHECK. Nothing is ever silently guessed: a field that is not in
the documents comes out as "NOT FOUND - verify manually".

Runs in Google Colab or on a local PC. No paid services.
"""

from __future__ import annotations

import hashlib
import html
import io
import json
import os
import re
import shutil
import sys
import time
import traceback
from dataclasses import dataclass, field as dc_field
from pathlib import Path

# Load .env file if it exists
_env_file = Path(__file__).parent / ".env"
if _env_file.exists():
    for line in _env_file.read_text().strip().split("\n"):
        if line and "=" in line and not line.startswith("#"):
            key, val = line.split("=", 1)
            os.environ[key.strip()] = val.strip()

# --------------------------------------------------------------------------
# CONFIG  (the notebook overrides these)
# --------------------------------------------------------------------------

CONFIG = {
    # Source of documents: "drive_link" or "local_folder"
    "source_mode": "drive_link",
    "drive_folder_url": "",
    "local_folder": "",

    # Working dirs
    "work_dir": "tender_work",
    "output_xlsx": "Tender_Summary.xlsx",
    "output_html": "Tender_Dashboard.html",

    # Folders/files to ignore (your own draft submissions, temp files)
    "exclude_dir_names": ["WORKING FOLDER", "SCAN_OUT", "__MACOSX"],
    "exclude_file_prefixes": ["~$", "."],

    # OCR
    "ocr_enabled": True,
    "ocr_dpi": 300,
    "ocr_lang": "eng",
    "ocr_max_pages_per_file": 60,      # safety valve on giant scans
    "text_layer_min_chars": 60,        # below this a page is treated as a scan

    # AI Engine (Groq 120B + Gemini Flash)
    "use_ai": True,
    "ai_provider": "auto",             # "auto", "groq", "gemini", or "both"
    "require_ai": True,                # Compulsory AI extraction

    # Groq API (High speed, 120B model)
    "use_groq": True,
    "groq_api_key": "",
    "groq_models": [
        "openai/gpt-oss-120b",
        "qwen/qwen3.8-27b",
        "groq/compound",
        "openai/gpt-oss-20b",
    ],

    # Gemini
    "use_gemini": True,
    "gemini_api_key": "",
    "gemini_models": [
        "gemini-3.6-flash",
        "gemini-3.7-flash",
        "gemini-3.5-flash-lite",
        "gemini-flash-latest",
    ],
    "gemini_chunk_chars": 200_000,     # per request; merged afterwards
    "gemini_max_chunks": 8,
    "gemini_retries": 4,               # rounds across the whole model list

    # Caching: skip re-processing tenders whose files have not changed
    "use_cache": True,
    "verbose": True,
}

NOT_FOUND = "NOT FOUND - verify manually"

FIELDS = [
    ("tender_name",      "1. Tender Name"),
    ("location",         "2. Location / Address"),
    ("purpose",          "3. Purpose / Audit Type"),
    ("period",           "4. Period"),
    ("estimated_cost",   "5. Tender Estimated Cost"),
    ("assignment_fees",  "6. Assignment Fees"),
    ("eligibility",      "7. Eligibility Criteria"),
    ("scope_of_work",    "8. Scope of Work"),
    ("penalty",          "9. Penalty"),
    ("emd",              "10. Tender EMD"),
    ("sd",               "11. Tender SD"),
    ("tender_fees",      "12. Tender Fees"),
    ("submission_date",  "13. Tender Submission Date"),
]

MONEY_FIELDS = {"estimated_cost", "assignment_fees", "emd", "sd", "tender_fees"}
DATE_FIELDS = {"submission_date"}
PROSE_FIELDS = {"eligibility", "scope_of_work", "penalty", "purpose"}


def log(*a):
    if CONFIG.get("verbose"):
        print(*a, flush=True)


# --------------------------------------------------------------------------
# 1. GOOGLE DRIVE  -  list a public folder recursively, download the files
# --------------------------------------------------------------------------

FOLDER_ID_RE = re.compile(r"/drive/folders/([\w-]+)")
FILE_ID_RE = re.compile(r"/file/d/([\w-]+)")
ANCHOR_RE = re.compile(r'<a[^>]+href="([^"]+)"[^>]*>(.*?)</a>', re.S | re.I)
TAG_RE = re.compile(r"<[^>]*>")


def drive_folder_id(url_or_id: str) -> str:
    url_or_id = (url_or_id or "").strip()
    m = FOLDER_ID_RE.search(url_or_id)
    if m:
        return m.group(1)
    m = re.search(r"[?&]id=([\w-]+)", url_or_id)
    if m:
        return m.group(1)
    return url_or_id.rstrip("/").split("/")[-1].split("?")[0]


def list_drive_folder(folder_id: str):
    """List one Drive folder via the public embeddedfolderview endpoint.

    Needs no API key and no login, as long as the folder is link-viewable.
    Returns [{'name','id','is_folder'}].
    """
    import requests

    url = f"https://drive.google.com/embeddedfolderview?id={folder_id}#list"
    r = requests.get(url, timeout=60,
                     headers={"User-Agent": "Mozilla/5.0 (tender-extractor)"})
    r.raise_for_status()
    items, seen = [], set()
    for href, inner in ANCHOR_RE.findall(r.text):
        name = html.unescape(TAG_RE.sub("", inner)).strip()
        if not name:
            continue
        fm, dm = FOLDER_ID_RE.search(href), FILE_ID_RE.search(href)
        if fm:
            fid, is_folder = fm.group(1), True
        elif dm:
            fid, is_folder = dm.group(1), False
        else:
            continue
        if fid in seen:
            continue
        seen.add(fid)
        items.append({"name": name, "id": fid, "is_folder": is_folder})
    return items


def _excluded_dir(name: str) -> bool:
    n = name.strip().upper()
    return any(n == x.strip().upper() for x in CONFIG["exclude_dir_names"])


def _excluded_file(name: str) -> bool:
    return any(name.startswith(p) for p in CONFIG["exclude_file_prefixes"])


def walk_drive(folder_id: str, rel: str = "", depth: int = 0, max_depth: int = 6):
    """Recursively yield (relative_path, file_id) for every downloadable file."""
    if depth > max_depth:
        return
    try:
        items = list_drive_folder(folder_id)
    except Exception as e:
        log(f"   ! could not list folder {folder_id}: {e}")
        return
    for it in items:
        if it["is_folder"]:
            if _excluded_dir(it["name"]):
                log(f"   - skipping (your own drafts): {rel}/{it['name']}")
                continue
            yield from walk_drive(it["id"], f"{rel}/{it['name']}".strip("/"),
                                  depth + 1, max_depth)
        else:
            if _excluded_file(it["name"]):
                continue
            yield (f"{rel}/{it['name']}".strip("/"), it["id"])


# rel path -> Drive file id, for the folder mirrored by the last download.
# The cache key folds these in, so replacing a document in Drive changes the
# key and forces that tender to be re-read.
DRIVE_IDS: dict[str, str] = {}

MANIFEST_NAME = ".drive_manifest.json"


def download_drive_folder(url_or_id: str, dest: Path) -> Path:
    """Mirror the shared Drive folder into `dest`, preserving structure.

    Skipping every file that already exists locally is what made the tool
    show yesterday's answers: correct a document in Drive and the old copy
    was kept forever. A manifest of rel-path -> Drive id is written beside
    the files, so a document whose id has changed is re-fetched and one that
    has been removed from Drive stops being read.
    """
    import gdown

    root_id = drive_folder_id(url_or_id)
    dest.mkdir(parents=True, exist_ok=True)
    files = list(walk_drive(root_id))
    if not files:
        raise RuntimeError(
            "No files found. Check that the Drive link is set to "
            "'Anyone with the link - Viewer', or switch source_mode to "
            "'local_folder' and use a mounted Drive path instead."
        )
    log(f"   found {len(files)} document(s) in Drive")

    manifest_path = dest / MANIFEST_NAME
    previous: dict[str, str] = {}
    if manifest_path.exists():
        try:
            previous = json.loads(manifest_path.read_text(encoding="utf-8"))
        except Exception:
            previous = {}

    current = {rel: fid for rel, fid in files}
    DRIVE_IDS.clear()
    DRIVE_IDS.update(current)

    fetched = replaced = 0
    for rel, fid in files:
        target = dest / rel
        target.parent.mkdir(parents=True, exist_ok=True)
        have = target.exists() and target.stat().st_size > 0
        changed = previous.get(rel) not in (None, fid)
        if have and not changed:
            continue
        if changed:
            log(f"   changed in Drive, re-downloading {rel}")
            replaced += 1
        else:
            log(f"   downloading {rel}")
            fetched += 1
        try:
            gdown.download(id=fid, output=str(target), quiet=True)
        except Exception as e:
            log(f"   ! download failed for {rel}: {e}")

    # Reconcile the whole mirror against what Drive actually lists now.
    # Trusting the manifest alone was not enough: the first run after any
    # upgrade has no manifest, so deleted tenders would survive on disk and
    # keep appearing in the summary. Every readable document under dest that
    # Drive no longer lists is dropped, whether or not we recorded it.
    wanted = {r.replace("\\", "/") for r in current}
    dropped = 0
    for path in sorted(dest.rglob("*")):
        if not path.is_file() or path.name == MANIFEST_NAME:
            continue
        if path.suffix.lower() not in READERS:
            continue
        rel = str(path.relative_to(dest)).replace("\\", "/")
        if rel in wanted:
            continue
        log(f"   removed from Drive, dropping local copy: {rel}")
        try:
            path.unlink()
            dropped += 1
        except OSError:
            pass

    # Prune directories the deletions emptied, deepest first, so a tender
    # folder that no longer exists in Drive stops being listed at all.
    for d in sorted((p for p in dest.rglob("*") if p.is_dir()),
                    key=lambda p: len(p.parts), reverse=True):
        try:
            next(d.iterdir())
        except StopIteration:
            try:
                d.rmdir()
            except OSError:
                pass
        except OSError:
            pass

    if replaced:
        log(f"   {replaced} document(s) had changed in Drive")
    if dropped:
        log(f"   {dropped} document(s) no longer in Drive were removed")
    try:
        manifest_path.write_text(json.dumps(current, indent=1),
                                 encoding="utf-8")
    except Exception:
        pass
    return dest


# --------------------------------------------------------------------------
# 2. TEXT EXTRACTION  -  PDF (with OCR fallback), DOCX, DOC, XLS/XLSX
# --------------------------------------------------------------------------

@dataclass
class Page:
    file: str          # relative filename
    page: int          # 1-based page / sheet number (0 = whole file)
    text: str
    ocr: bool = False
    boilerplate: bool = False   # standing contract conditions, not this bid

    @property
    def ref(self) -> str:
        p = f"p.{self.page}" if self.page else "whole file"
        return f"{self.file} ({p}{', OCR' if self.ocr else ''})"


def _clean(t: str) -> str:
    t = t.replace("\u00a0", " ").replace("\r", "\n")
    t = re.sub(r"[ \t]+", " ", t)
    t = re.sub(r"\n{3,}", "\n\n", t)
    return t.strip()


def read_pdf(path: Path) -> list[Page]:
    import fitz  # PyMuPDF

    pages: list[Page] = []
    try:
        doc = fitz.open(str(path))
    except Exception as e:
        log(f"   ! cannot open PDF {path.name}: {e}")
        return pages

    ocr_used = 0
    for i, pg in enumerate(doc, start=1):
        try:
            txt = pg.get_text("text") or ""
        except Exception:
            txt = ""
        is_ocr = False
        if (len(txt.strip()) < CONFIG["text_layer_min_chars"]
                and CONFIG["ocr_enabled"]
                and ocr_used < CONFIG["ocr_max_pages_per_file"]):
            ocr_txt = _ocr_page(pg)
            if len(ocr_txt.strip()) > len(txt.strip()):
                txt, is_ocr = ocr_txt, True
                ocr_used += 1
        txt = _clean(txt)
        if txt:
            pages.append(Page(path.name, i, txt, is_ocr))
    doc.close()
    if ocr_used:
        log(f"     ({ocr_used} scanned page(s) OCR-ed in {path.name})")
    return pages


_OCR_WARNED = False


def _ocr_page(pg) -> str:
    global _OCR_WARNED
    try:
        # pyrefly: ignore [missing-import]
        import pytesseract
        from PIL import Image
        import fitz

        zoom = CONFIG["ocr_dpi"] / 72.0
        pix = pg.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img, lang=CONFIG["ocr_lang"])
    except ImportError:
        if not _OCR_WARNED:
            log("   - Tesseract OCR module not installed (install pytesseract for scanned page OCR)")
            _OCR_WARNED = True
        return ""
    except Exception as e:
        if not _OCR_WARNED:
            log(f"   ! OCR unavailable/failed: {e}")
            _OCR_WARNED = True
        return ""


def read_docx(path: Path) -> list[Page]:
    try:
        import docx
    except Exception as e:
        log(f"   ! python-docx missing: {e}")
        return []
    try:
        d = docx.Document(str(path))
    except Exception as e:
        log(f"   ! cannot open {path.name}: {e}")
        return []
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    txt = _clean("\n".join(parts))
    return [Page(path.name, 0, txt)] if txt else []


def read_doc(path: Path) -> list[Page]:
    """Legacy .doc: try LibreOffice conversion, else report as unreadable."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        log(f"   ! {path.name}: legacy .doc needs LibreOffice; skipped")
        return []
    outdir = path.parent / "_conv"
    outdir.mkdir(exist_ok=True)
    os.system(f'"{soffice}" --headless --convert-to docx --outdir '
              f'"{outdir}" "{path}" >/dev/null 2>&1')
    conv = outdir / (path.stem + ".docx")
    return read_docx(conv) if conv.exists() else []


def read_excel(path: Path) -> list[Page]:
    try:
        import pandas as pd
    except Exception:
        return []
    pages = []
    try:
        sheets = pd.read_excel(str(path), sheet_name=None, header=None,
                               dtype=str)
    except Exception as e:
        log(f"   ! cannot read {path.name}: {e}")
        return []
    for n, (name, df) in enumerate(sheets.items(), start=1):
        df = df.fillna("")
        rows = ["\t".join(str(v) for v in r if str(v).strip())
                for r in df.values.tolist()]
        txt = _clean(f"[Sheet: {name}]\n" + "\n".join(r for r in rows if r.strip()))
        if len(txt) > 40:
            pages.append(Page(path.name, n, txt))
    return pages


READERS = {
    ".pdf": read_pdf, ".docx": read_docx, ".doc": read_doc,
    ".xls": read_excel, ".xlsx": read_excel, ".xlsm": read_excel,
    ".csv": read_excel, ".txt": lambda p: [Page(p.name, 0,
                                                _clean(p.read_text(errors="ignore")))],
}


def read_any(path: Path) -> list[Page]:
    fn = READERS.get(path.suffix.lower())
    if not fn:
        log(f"   - unsupported file type, skipped: {path.name}")
        return []
    try:
        return fn(path)
    except Exception:
        log(f"   ! failed reading {path.name}\n{traceback.format_exc(limit=2)}")
        return []


# Documents most likely to hold the summary facts get read/ranked first.
PRIORITY_HINTS = [
    ("gem-bidding", 100), ("gem_bidding", 100), ("bid document", 90),
    ("tender-summary", 95), ("tender summary", 95), ("nib", 85), ("nit", 85),
    ("commercial-buyer", 80), ("tender", 60), ("boq", 40), ("attachment", 20),
    ("upload documents", 5),
]


def file_priority(name: str) -> int:
    n = name.lower()
    return max((w for k, w in PRIORITY_HINTS if k in n), default=50)


# A tender pack usually includes the department's standing contract
# conditions - "Instructions to Bidders", "General Conditions of Contract" -
# reissued unchanged with every tender it floats. They describe the
# department's business in general, not this assignment, so every field taken
# from them is wrong in a way that reads as confident and specific: NTPC's
# registered office in New Delhi instead of the audit site in Maharashtra,
# "construction, erection and commissioning of power projects" instead of a
# statutory audit, a fatal-injury penalty instead of an audit penalty.
#
# They are also long, and named things like TENDER.pdf that score highly on
# file name alone, so they were being read first and winning every field.
BOILERPLATE_RE = re.compile(
    r"INSTRUCTIONS?\s+TO\s+BIDDERS|GENERAL\s+CONDITIONS\s+OF\s+CONTRACT|"
    r"TABLE\s+OF\s+CLAUSES|CONDITIONS\s+OF\s+CONTRACT\s+FOR|"
    r"STANDARD\s+(?:BIDDING|TENDER)\s+DOCUMENT", re.I)


def mark_boilerplate(pages: list[Page]) -> int:
    """Flag files that are the department's standing conditions, not this bid.

    Judged on the opening pages: a genuine tender may cite these documents in
    passing, but only the template itself leads with them.
    """
    by_file: dict[str, list[Page]] = {}
    for p in pages:
        by_file.setdefault(p.file, []).append(p)

    flagged = 0
    for fname, fpages in by_file.items():
        head = sorted(fpages, key=lambda p: p.page)[:3]
        hits = sum(1 for p in head if BOILERPLATE_RE.search(p.text or ""))
        # One mention could be a reference; leading with it twice, or in a
        # document this long, means the file *is* the standing conditions.
        if hits >= 2 or (hits and len(fpages) >= 50):
            for p in fpages:
                p.boilerplate = True
            flagged += 1
            log(f"   note: {fname} looks like standard contract conditions, "
                f"not this tender - used only to fill gaps")
    return flagged


# --------------------------------------------------------------------------
# 3. RULES LAYER  -  deterministic label / regex / section extraction
# --------------------------------------------------------------------------

@dataclass
class Cand:
    value: str = ""
    ref: str = ""
    conf: str = "low"          # high | medium | low
    numeric: float | None = None
    raw: str = ""


CUR = r"(?:Rs\.?|INR|\u20b9|Rupees)"
NUM = r"\d[\d,]*(?:\.\d+)?"
MULT = r"(?:lakh?s?|lacs?|lac|crores?|cr\.?|thousand|million|mn)"

AMOUNT_RE = re.compile(rf"({CUR})?\s*({NUM})\s*(/-)?\s*({MULT})?", re.I)

# Labels from structured key-value templates (GeM and the like) where a bare
# number with no "Rs." really is the amount. Everywhere else a bare number is
# rejected, because in running prose it is nearly always a year, a clause
# number, a tender number or a file name.
BARE_OK_LABELS = {
    r"EMD\s*Amount",
    r"Estimated\s*Bid\s*Value",
    r"Estimated\s*(?:Cost|Value|Amount|Contract\s*Value)",
}
BARE_MIN = 1000          # a bare number below this is not a rupee amount
BARE_MAX = 1e10          # above this it is an account/reference number
FILENAME_CTX = re.compile(r"\.pdf|\.xls|\.doc|Bidding-|Document-", re.I)

MULTIPLIERS = {
    "lakh": 1e5, "lakhs": 1e5, "lakh?s": 1e5, "lac": 1e5, "lacs": 1e5,
    "crore": 1e7, "crores": 1e7, "cr": 1e7,
    "thousand": 1e3, "million": 1e6, "mn": 1e6,
}

DATE_RE = re.compile(
    r"(\d{1,2}[-/.\s](?:\d{1,2}|Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec"
    r"|January|February|March|April|June|July|August|September|October|November"
    r"|December)[a-z]*[-/.\s]\d{2,4}(?:\s+\d{1,2}:\d{2}(?::\d{2})?)?)"
    r"|((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})",
    re.I)


def parse_amount(window: str, allow_bare: bool = False):
    """Return (numeric_value, display_string) for the first real amount found.

    A number is only accepted as money when it carries a currency marker
    ("Rs.", "INR", the rupee sign), a "/-" suffix, or a multiplier word
    ("lakh", "crore"). Bare digits are accepted only when `allow_bare` is set,
    which happens for structured key-value labels such as GeM's "EMD Amount".
    This is what stops years, clause numbers, tender numbers and account
    numbers being reported as rupee figures.
    """
    for m in AMOUNT_RE.finditer(window):
        cur, raw_num, slash = m.group(1), m.group(2), m.group(3)
        mult = (m.group(4) or "").lower().strip().rstrip(".")
        if not raw_num or not raw_num.strip(",."):
            continue

        digits = raw_num.replace(",", "").split(".")[0]
        # Leading zeros mean an identifier (account no., reference no.), not money.
        if len(digits) > 1 and digits.startswith("0"):
            continue
        if len(digits) > 12:
            continue

        # Embedded in an identifier such as "T-168" or "Bidding-9722687"?
        s = m.start(2)
        if s > 0 and (window[s - 1] in "-/.:#" or window[s - 1].isalnum()):
            continue

        try:
            val = float(raw_num.replace(",", ""))
        except ValueError:
            continue
        if mult:
            val *= MULTIPLIERS.get(mult, 1)

        if window[m.end():m.end() + 2].lstrip().startswith("%"):
            continue

        has_marker = bool(cur or slash or mult)
        if not has_marker:
            if not allow_bare:
                continue
            if val < BARE_MIN or val > BARE_MAX:
                continue
            # A bare 4-digit number in the 1900-2100 range is a year.
            if 1900 <= val <= 2100 and float(val).is_integer():
                continue
            # A bare number sitting next to a file name is part of the file name.
            if FILENAME_CTX.search(window[max(0, s - 60): m.end() + 30]):
                continue

        disp = f"Rs. {val:,.0f}" if val == int(val) else f"Rs. {val:,.2f}"
        return val, f"{disp}  [as written: {m.group(0).strip()}]"
    return None, ""


def parse_date(window: str) -> str:
    m = DATE_RE.search(window)
    return m.group(0).strip() if m else ""


# --- accuracy patch 2026-08-26 ---

# "Not Applicable" against EMD is an ANSWER, not a miss. Reporting it as
# NOT FOUND sends the reader hunting for something the tender already
# settled. Seen in three separate documents of tender 2026-9533.
EXPLICIT_NIL_RE = re.compile(
    r"\b(?:not\s*applicable|not\s*required|nil|none|n\.?a\.?|exempt(?:ed)?|"
    r"no\s+emd|without\s+emd|waived)\b", re.I)

# Words that turn a nearby rupee figure into a THRESHOLD or a CEILING
# rather than the amount payable. Real culprit in 2026-9533:
#   "Security Deposit amount up to Rs. 1,00,000/- must be submitted
#    through Electronic Fund Transfer (EFT) only"
# — a payment-mode rule, read as the deposit itself.
MONEY_QUALIFIER_RE = re.compile(
    r"\b(?:up\s*to|upto|not\s*exceeding|exceeding|in\s*excess\s*of|"
    r"more\s+than|less\s+than|below|above|at\s*least|minimum\s+of|"
    r"maximum\s+of|whichever|threshold|per\s*day|per\s*week|per\s*month|"
    r"per\s*visit|per\s*branch|per\s*sitting|slab)\b", re.I)

# Security deposit / performance security in these tenders is stated as a
# percentage of contract value, never as a rupee figure. Capturing the
# percentage is both correct and more useful than a derived amount.
PCT_OF_RE = re.compile(
    r"(\d{1,2}(?:\.\d+)?)\s*(?:%|per\s*cent|percent)\s*(?:of|on)\s+"
    r"((?:the\s+)?(?:total\s+)?(?:awarded\s+)?"
    r"(?:contract|order|bid|tender|work|purchase\s*order)\s*"
    r"(?:price|value|amount|cost))", re.I)

# PyMuPDF leaves an unmapped glyph as U+FFFD. The NTPC GCC states the
# liquidated-damages rate with a vulgar fraction, so the rate extracts as
# "damages@ \ufffd percent per week". Emitting that silently would be worse
# than saying the character could not be read.
UNREADABLE_GLYPH = "\ufffd"


def parse_percent_of(window: str) -> str:
    """Return e.g. "5.0% of the total Contract Price", or ""."""
    m = PCT_OF_RE.search(window)
    if not m:
        return ""
    basis = re.sub(r"\s+", " ", m.group(2)).strip()
    return "%s%% of %s" % (m.group(1), basis)


def parse_amount_at(window: str, allow_bare: bool = False):
    """parse_amount, but also returns where in the window the figure sat.

    The offset is what lets the caller check whether a qualifier such as
    "up to" appears between the label and the number.
    """
    for m in AMOUNT_RE.finditer(window):
        val, disp = parse_amount(m.group(0), allow_bare=allow_bare)
        if val is not None:
            return val, disp, m.start()
    return None, "", -1


# field -> label regexes. Order matters: most specific label first.
LABELS: dict[str, list[str]] = {
    "emd": [
        r"EMD\s*Amount",
        r"Earnest\s*Money\s*Deposit",
        r"\bEMD\b",
        r"Earnest\s*Money",
        r"Bid\s*Security(?:\s*Amount)?",
    ],
    "tender_fees": [
        r"Tender\s*(?:Document\s*)?Fee",
        r"Cost\s*of\s*(?:the\s*)?(?:Tender|Bid|RFP)\s*Document",
        r"Bid\s*Document\s*(?:Fee|Cost)",
        r"(?:e-?)?Tender\s*Processing\s*Fee",
        r"Document\s*Fee",
        r"Application\s*Fee",
    ],
    "sd": [
        r"Security\s*Deposit",
        r"Performance\s*(?:Bank\s*)?Guarantee",
        r"Performance\s*Security",
        r"\bPBG\b",
        r"Retention\s*Money",
    ],
    "estimated_cost": [
        r"Estimated\s*Bid\s*Value",
        r"Estimated\s*(?:Cost|Value|Amount|Contract\s*Value)",
        r"Approx(?:imate)?\s*(?:Cost|Value)",
        r"Tender\s*Value",
        r"Contract\s*Value",
        r"Value\s*of\s*(?:the\s*)?(?:Work|Contract|Tender)",
    ],
    "assignment_fees": [
        r"Assignment\s*Fee",
        r"(?:Audit|Professional|Consultancy)\s*Fee",
        r"Remuneration",
        r"Fee\s*(?:Payable|Quoted)",
    ],
    "submission_date": [
        r"Bid\s*End\s*Date(?:\s*/\s*Time)?",
        r"Last\s*Date\s*(?:and|&)?\s*Time\s*(?:for|of)\s*(?:Online\s*)?"
        r"(?:Submission|Receipt|Uploading)",
        r"Last\s*Date\s*(?:for|of)\s*Submission",
        r"(?:Bid|Tender)\s*Submission\s*(?:End|Closing|Last)\s*Date",
        r"Due\s*Date\s*(?:and|&)?\s*Time",
        r"Closing\s*Date(?:\s*(?:and|&)\s*Time)?",
    ],
    "period": [
        r"Contract\s*Period",
        r"Period\s*of\s*(?:Contract|Audit|Engagement|Assignment|Service)",
        r"Audit\s*Period",
        r"Duration\s*of\s*(?:Contract|Work|Assignment)",
        r"for\s*the\s*(?:Financial\s*)?[Yy]ear",
        r"Bid\s*Offer\s*Validity",
    ],
    "location": [
        # Site/work address first. "Organisation Name" used to win and
        # returned "Ntpc Limited" for a tender whose site address was
        # spelled out three times under these labels.
        r"Location\s*of\s*Work",
        r"Site\s*location",
        r"Delivery\s*Address",
        r"Place\s*of\s*(?:Work|Audit|Posting)",
        r"Address\s*of\s*(?:the\s*)?(?:Office|Department|Organisation)",
        r"Buyer\s*Name\s*/?\s*Address",
        # Full form only. Matching bare "Consignee" hits the plural heading
        # "Consignees/Reporting Officer and Quantity" and captures its leftovers.
        r"Consignee\s*/\s*Reporting\s*Officer\s*/\s*Address",
        r"\bLocation\b",
        r"Delivery\s*district",
        # Names of the buying entity, not a place. Last resort only.
        r"Office\s*Name",
        r"Organisation\s*Name",
    ],
    "purpose": [
        # "Item Category" on a GeM custom bid is always the useless
        # "Custom Bid for Services". The work name carries the real answer.
        r"Name\s*of\s*(?:the\s*)?Work",
        r"Type\s*of\s*Audit",
        r"Nature\s*of\s*(?:Work|Service|Assignment)",
        r"Nature\s*of\s*Requirement",
        r"Products?",
        r"Subject",
        r"Item\s*Category",
    ],
    "tender_name": [
        r"Bid\s*Number",
        r"(?:Tender|NIT|NIB|RFP)\s*(?:Reference\s*)?(?:No|Number|ID)",
        r"Name\s*of\s*(?:the\s*)?(?:Work|Tender|Project)",
    ],
}

SECTION_HEADINGS: dict[str, list[str]] = {
    "scope_of_work": [
        r"SCOPE\s*OF\s*(?:THE\s*)?WORK",
        r"SCOPE\s*OF\s*(?:AUDIT|SERVICES?|ASSIGNMENT)",
        r"TERMS\s*OF\s*REFERENCE",
        r"DETAILED\s*SCOPE",
        r"WORK\s*TO\s*BE\s*(?:DONE|PERFORMED)",
        r"DUTIES\s*(?:AND|&)\s*RESPONSIBILITIES",
    ],
    "eligibility": [
        r"ELIGIBILITY\s*CRITERIA",
        r"PRE[\s-]*QUALIFICATION(?:\s*CRITERIA)?",
        r"MINIMUM\s*ELIGIBILITY",
        r"QUALIFICATION\s*(?:CRITERIA|REQUIREMENTS)",
        r"WHO\s*CAN\s*APPLY",
        r"ELIGIBILITY\s*(?:CONDITIONS|REQUIREMENTS)",
    ],
    "penalty": [
        r"PENALT(?:Y|IES)",
        r"LIQUIDATED\s*DAMAGES?",
        r"PENAL\s*(?:CLAUSE|PROVISION)",
        r"DEDUCTIONS?\s*(?:AND|&)\s*PENALT",
    ],
}

# A line that looks like the start of the next clause / heading.
NEXT_HEADING_RE = re.compile(
    r"^\s*(?:(?:\d+(?:\.\d+)*)[).]?\s+[A-Z]"
    r"|[A-Z][A-Z0-9 ,\-&/()'.]{8,}\s*:?\s*$"
    r"|(?:ANNEXURE|APPENDIX|SECTION|CHAPTER|CLAUSE|PART)\b)")

SUBITEM_RE = re.compile(r"^\s*\(?[a-z0-9ivx]{1,3}[).]")

# For eligibility only: the GeM criteria table numbers its own
# sub-questions, so a numbered line is content, not the next heading.
# Stop only at a genuine all-caps heading or a document division.
ELIG_BREAK_RE = re.compile(
    r"^\s*(?:[A-Z][A-Z0-9 ,\-&/()'.]{14,}\s*:?\s*$"
    r"|(?:ANNEXURE|APPENDIX|SECTION|CHAPTER|PART|CHECKLIST|DECLARATION)\b)")


def _ordered_pages(pages: list[Page]) -> list[Page]:
    """Read the documents most likely to carry the facts first.

    Standing contract conditions go last whatever their file name scores:
    TENDER.pdf matches the "tender" hint and outranked the real documents,
    which is how a 438-page O&M template came to supply seven of thirteen
    fields at high confidence.
    """
    return sorted(pages, key=lambda p: (p.boilerplate,
                                        -file_priority(p.file), p.file, p.page))


class _Block:
    """One file's pages joined into a single searchable text.

    Eligibility criteria routinely run over several pages; capturing per
    page cut tender 2026-9533's criteria off after the first of four,
    losing the CPCB-authorised-auditor requirement that decides whether
    the firm may bid at all. `ref` reports the page the match started on.
    """

    def __init__(self, pages: list[Page]):
        pages = sorted(pages, key=lambda p: p.page)
        self.file = pages[0].file
        self._starts: list[tuple[int, Page]] = []
        chunks, pos = [], 0
        for p in pages:
            self._starts.append((pos, p))
            chunks.append(p.text)
            pos += len(p.text) + 1
        self.text = "\n".join(chunks)
        self._match_at = 0

    @property
    def ref(self) -> str:
        page = self._starts[0][1]
        for start, p in self._starts:
            if start <= self._match_at:
                page = p
            else:
                break
        return page.ref


def _file_blocks(pages: list[Page]) -> list[_Block]:
    by_file: dict[str, list[Page]] = {}
    for p in pages:
        by_file.setdefault(p.file, []).append(p)
    blocks = [_Block(v) for v in by_file.values()]
    return sorted(blocks, key=lambda b: (-file_priority(b.file), b.file))


HEADINGISH_RE = re.compile(
    r"^\s*(?:\d+(?:\.\d+)*[).]?\s*)?[A-Z][A-Z0-9 ,\-&/()'.]{4,}\s*:?\s*$")

# A line ending like this is mid-sentence, so the value continues on the next line.
CONTINUES_RE = re.compile(
    r"(?:\b(?:of|for|the|and|to|in|by|with|at|on|from|a|an)\b|[,;:\-\u2013])\s*$", re.I)


DEVANAGARI_RE = re.compile(r"[ऀ-ॿ]")


def _is_junk_line(s: str) -> bool:
    """Bilingual-template noise or an unfilled fill-in-the-blank line.

    GeM's forms echo every label in Hindi beside the English one, and an
    unanswered field ("N/a") often sits glued to that Hindi echo rather than
    to a real value. A blank field on a print template shows as a run of
    dots/underscores/dashes ("......, dated: ...... for Providing ......").
    Both look like a value if only length is checked, so are filtered here.
    """
    if not s:
        return True
    if DEVANAGARI_RE.search(s):
        return True
    alnum = sum(ch.isalnum() for ch in s)
    return len(s) >= 15 and alnum / len(s) < 0.35


def _text_value(window: str) -> str:
    """Pull a label's value, following it onto the next line when it wraps."""
    lines = [ln.strip() for ln in window.split("\n")]
    idx = next((i for i, ln in enumerate(lines) if ln.strip(" :|-.")), None)
    if idx is None:
        return ""
    first = lines[idx].strip(" :|-")
    # If we landed on the next heading instead of a value, this is not it.
    if HEADINGISH_RE.match(first) and len(first) < 60:
        return ""
    # We landed inside the label itself, e.g. matching "Consignee" inside
    # "Consignees/Reporting Officer" and capturing the leftover "s/Reporting...".
    if re.match(r"^[a-z]{1,3}\s*[/)\]]", first):
        return ""
    # The label was a word inside a sentence, not a field, so what follows is
    # the rest of that clause. "...at the tender opening location, before the
    # deadline for submission of bids..." was reported as tender 2026-9533's
    # address. A real value never opens mid-clause.
    if re.match(r"^[,;]", first):
        return ""
    if _is_junk_line(first):
        return ""
    parts = [first]
    for nxt in lines[idx + 1: idx + 4]:
        if not nxt or HEADINGISH_RE.match(nxt) or _is_junk_line(nxt):
            break
        if not (CONTINUES_RE.search(parts[-1]) or nxt[:1].islower()):
            break
        parts.append(nxt)
        if len(" ".join(parts)) > 300:
            break
    return re.sub(r"\s{2,}", " ", " ".join(parts)).strip(" :|-")


# Values that are technically present but say nothing. Accepting one of
# these stops the search and hides the real answer further down the page.
GENERIC_VALUES = {
    "purpose": re.compile(
        r"^(?:custom\s*bid(?:\s*for\s*services?)?|services?|goods?|works?|"
        r"n\.?a\.?|not\s*applicable|others?|miscellaneous)\s*$", re.I),
    "location": re.compile(
        r"^(?:n\.?a\.?|not\s*applicable|india|as\s*per\s*.*|"
        r"[a-z0-9 .&-]*(?:limited|ltd\.?|corporation|corp\.?|"
        r"nigam|board|authority|company)\s*)$", re.I),
}

# --- accuracy patch round 2 ---
# Words that are the NEXT label on a blank proforma, not a value. Tender
# documents are full of "Name of Work: ..........  Date: .........."; the
# dotted blank is skipped as junk and the following label gets picked up.
FORM_LABEL_RE = re.compile(
    r"^(?:date|dated|signature|sign|name|address|place|seal|stamp|"
    r"designation|witness|to|from|ref|reference|subject|sub|page|"
    r"sr\.?\s*no\.?|s\.?\s*no\.?|annexure|appendix|amount|total|nil)"
    r"\s*:?\s*$", re.I)

# An address is expected to carry a PIN code or several comma-separated
# parts. A single proper noun is a name, not a place.
ADDRESSISH_RE = re.compile(r"\b\d{6}\b|,.*,|\bP\.?O\.?\b|\bDist\b|\bPIN\b", re.I)


def _value_is_weak(field: str, value: str) -> bool:
    """True when a value parsed cleanly but carries no information."""
    v = value.strip()
    rx = GENERIC_VALUES.get(field)
    if rx and rx.match(v):
        return True
    if field in ("purpose", "location") and FORM_LABEL_RE.match(v):
        return True
    if field == "purpose" and len(re.findall(r"[A-Za-z]{2,}", v)) < 3:
        # "Date", "Audit", "Services" - a label or a bare category.
        return True
    if field == "location" and len(v) < 60 and not ADDRESSISH_RE.search(v):
        return True
    return False


def _line_anchored(text: str, start: int, end: int) -> bool:
    """True if the label starts its own line, or is immediately followed by ':'."""
    # A label followed by a comma is a word in a sentence, not a field. PDF
    # extraction wraps lines mid-sentence, so "at the tender opening\nlocation,
    # before the deadline..." puts a bare "location" at the start of a line and
    # looks exactly like a heading; that fragment was reported as tender
    # 2026-9533's address.
    after = text[end:end + 1]
    if after == ",":
        return False
    bol = text.rfind("\n", 0, start) + 1
    if re.fullmatch(r"[\s\u2022*\-]*(?:\d+(?:\.\d+)*[).]?\s*)?", text[bol:start]):
        # Only a genuine line start counts. If the previous line runs on
        # without terminal punctuation, this is a wrapped sentence.
        prev_end = bol - 1
        if prev_end > 0:
            prev = text.rfind("\n", 0, prev_end) + 1
            prev_line = text[prev:prev_end].rstrip()
            if prev_line and prev_line[-1] not in ".:;?!)]" and \
                    text[start:end].islower():
                return False
        return True
    return ":" in text[end:end + 3]


def _label_lookup(pages: list[Page], patterns: list[str], kind: str,
                  field: str = "") -> Cand:
    # Pass 1 trusts only labels that head a line or carry a colon - that is a
    # real field. Pass 2 relaxes it, so a label buried in prose is a fallback,
    # never a first choice.
    weak_money: list[Cand] = []
    weak_text: list[Cand] = []
    for strict in (True, False):
        for pat in patterns:
            rx = re.compile(pat, re.I)
            for pg in _ordered_pages(pages):
                for m in rx.finditer(pg.text):
                    if strict and not _line_anchored(pg.text, m.start(), m.end()):
                        continue
                    # Money gets a short window. Many portals dump a page as
                    # one short table cell per line with no real structure, so
                    # a wide window drifts clean past an unrelated label (e.g.
                    # "EMD BG ... to be uploaded on GeM Portal" has no amount
                    # at all) into a completely different field's number
                    # 200+ characters later. A real amount sits right next to
                    # its label; if it isn't within a short distance, it is
                    # not this field's value.
                    span = 140 if kind == "money" else 320
                    window = pg.text[m.end(): m.end() + span]
                    # drop the bilingual Hindi echo / separators on the label
                    window = re.sub(r"^[ \t:/|.\-\u2013\u2014]*", "", window)
                    conf = "high" if strict else "medium"
                    if kind == "money":
                        # A percentage of contract value is how security
                        # deposits and performance guarantees are actually
                        # specified. Prefer it over any rupee figure nearby.
                        pct = parse_percent_of(window)
                        if pct:
                            return Cand(pct, pg.ref, conf, None, m.group(0))

                        # Bare digits only from a structured label, matched
                        # line-anchored. In prose, currency must be present.
                        val, disp, at = parse_amount_at(
                            window, allow_bare=strict and pat in BARE_OK_LABELS)
                        if val is not None:
                            # "...deposit up to Rs 1,00,000 must be paid by
                            # EFT" states a threshold, not the deposit.
                            if MONEY_QUALIFIER_RE.search(window[:at]):
                                weak_money.append(
                                    Cand(disp, pg.ref, "low", val, m.group(0)))
                            else:
                                return Cand(disp, pg.ref, conf, val, m.group(0))
                        elif EXPLICIT_NIL_RE.search(window[:120]):
                            # Stated as nil/not applicable: a real answer.
                            return Cand("Not Applicable / Nil (as stated)",
                                        pg.ref, conf, 0.0, m.group(0))
                    elif kind == "date":
                        d = parse_date(window)
                        if d:
                            return Cand(d, pg.ref, conf, None, m.group(0))
                    else:
                        v = _text_value(window)
                        if len(v) >= 3:
                            c = Cand(v[:400], pg.ref,
                                     "medium" if strict else "low",
                                     None, m.group(0))
                            # Keep a placeholder value in reserve, but go on
                            # looking for one that actually says something.
                            if _value_is_weak(field, v):
                                weak_text.append(c)
                            else:
                                return c
    if weak_money:
        c = weak_money[0]
        return Cand(c.value + "   [context suggests a threshold, not the "
                    "amount payable - verify]", c.ref, "low", c.numeric, c.raw)
    if weak_text:
        return weak_text[0]
    return Cand()


def _capture_section(pages: list[Page], patterns: list[str],
                     max_chars: int = 4000,
                     min_before_break: int = 200,
                     span_pages: bool = False,
                     break_re: "re.Pattern | None" = None) -> Cand:
    """Capture from a heading down to the next heading-looking line."""
    best = Cand()
    units = _file_blocks(pages) if span_pages else _ordered_pages(pages)
    for pat in patterns:
        rx = re.compile(pat, re.I)
        for pg in units:
            # Only a heading that starts its own line is a real section
            # heading. The same words inside a sentence ("...as defined in the
            # scope of work above...") are a reference, not a section.
            m = next((mm for mm in rx.finditer(pg.text)
                      if _line_anchored(pg.text, mm.start(), mm.end())), None)
            if not m:
                continue
            if isinstance(pg, _Block):
                pg._match_at = m.start()
            lines, first = [], True
            for ln in pg.text[m.end():].split("\n"):
                if first:
                    first = False
                    if ln.strip(" :.-"):
                        lines.append(ln.strip())
                    continue
                joined_len = len("\n".join(lines))
                stop = break_re or NEXT_HEADING_RE
                if (joined_len > min_before_break and stop.match(ln)
                        and not SUBITEM_RE.match(ln)):
                    break
                lines.append(ln.rstrip())
                if joined_len > max_chars:
                    break
            body = _clean("\n".join(lines))
            # A section that opens mid-sentence means the heading match landed
            # inside a paragraph, so the text is a fragment, not the section.
            if body and (body[0].islower() or body[0] in ",;/)]}.-–"):
                continue
            if len(body) > len(best.value):
                best = Cand(body[:max_chars], pg.ref,
                            "high" if len(body) > 300 else "medium")
    return best


# The sentence that actually states the liquidated-damages rate, e.g.
#   "liable for payment of liquidated damages@ 1/2 percent per week, not as
#    penalty, on the Contract Value ... subject to a maximum of 5% of the
#    Contract Value"
# Capturing this beats capturing the clause heading, which in 2026-9533
# said only "Liquidated Damages / As per GCC" and then bled into the next
# two clauses about site inspection.
LD_RATE_RE = re.compile(
    r"(?:liquidated\s+damages?|penalt(?:y|ies))\b[^.]{0,400}?"
    r"(?:\d{1,2}(?:\.\d+)?\s*(?:%|per\s*cent|percent)|\ufffd|\u00bd|\u00bc|\u00be)"
    r"[^.]{0,400}?\.", re.I | re.S)

LD_CAP_RE = re.compile(
    r"maximum\s+of\s+(\d{1,2}(?:\.\d+)?)\s*(?:%|per\s*cent|percent)\s*"
    r"of\s+the\s+([A-Za-z ]{0,40}?(?:contract|order)\s*(?:value|price))", re.I)


RATE_PER_RE = re.compile(r"per\s*(?:week|day|month|fortnight)", re.I)


def _score_penalty(body: str, has_cap: bool) -> int:
    """Rank rate sentences. The operative clause states a rate AND a cap.

    Needed because a tender mentions penalties in several places -
    banning policy, non-performance, fraud - and only one of them is the
    delay penalty the bidder is exposed to. In 2026-9533 the first match
    in page order was a 10% banning-policy clause on p.42, while the
    operative liquidated damages sat in GCC 25.5.1 on p.90.
    """
    low = body.lower()
    score = 0
    if "liquidated damage" in low:
        score += 4
    if RATE_PER_RE.search(low):
        score += 3
    if has_cap or "maximum of" in low:
        score += 2
    if "contract value" in low or "contract price" in low:
        score += 1
    # A clause about forfeiting or banning is not the delay penalty.
    if any(w in low for w in ("banning", "blacklist", "debar", "forfeit")):
        score -= 3
    return score


def find_penalty_rate(pages: list[Page]) -> Cand:
    """The best liquidated-damages / penalty rate sentence in the tender."""
    best, best_score = Cand(), 0
    for pg in _ordered_pages(pages):
        for m in LD_RATE_RE.finditer(pg.text):
            body = _clean(m.group(0)).strip()
            if len(body) < 40:
                continue
            cap = LD_CAP_RE.search(pg.text[m.start(): m.start() + 1200])
            score = _score_penalty(body, bool(cap))
            if score <= 0 or score <= best_score:
                continue
            if cap and "maximum" not in body.lower():
                body += ("  (capped at %s%% of the %s)"
                         % (cap.group(1), cap.group(2)))
            if UNREADABLE_GLYPH in body:
                # The rate is a vulgar fraction the PDF font does not map.
                # Say so rather than shipping a mystery character.
                best = Cand(body.replace(UNREADABLE_GLYPH, "[?]")
                            + "   [the rate character did not extract - "
                            "read it off the source page]", pg.ref, "low")
            else:
                best = Cand(body, pg.ref, "high")
            best_score = score
    return best


GEM_BID_NO_RE = re.compile(r"(GEM/\d{4}/[BR]/\d+)", re.I)
EPBG_RE = re.compile(r"ePBG\s*Percentage\s*\(?%?\)?\s*[:\-]?\s*(\d+(?:\.\d+)?)", re.I)


# --- Tender247 "AI Generated Tender Summary" sheets -----------------------
#
# Most folders ship one of these alongside the bid documents. It is a plain
# two-column table that already states the cost, the EMD, the deadline, the
# site and the eligibility - the very fields that are hardest to find in a
# 500-page bid pack. Reading it directly is worth more than any amount of
# regex tuning against the long documents.
#
# The labels wrap across lines in the extracted text ("Bid End Date\nTime
# 25-08-2026"), so the page is flattened to one line before matching and
# each value is taken as whatever sits between its label and the next one.

SUMMARY_LABELS = [
    "Tender Id", "GEM Bid number", "Bid End Date Time", "Bid Opening Date Time",
    "Bid Offer Validity From EndDate", "Ministry State Name", "Department Name",
    "Organisation Name", "Office Name", "Item Category", "Contract Period",
    "Site location", "Bid to Ra Enabled", "Type of Bid",
    "Time Allowed for Clarifications", "Evaluation Method",
    "Emd Instrument Type", "Completion Period", "MSE Purchase Preference",
    "Category", "Delivery district", "Location", "Products", "State",
    "Estimated Cost", "EMD Value", "Emd Mode Type", "Mediation Clause",
    "Arbitration Clause", "Document required from seller", "Pre Bid Meeting",
    "Last date for Seeking Clarification", "Joint Venture OR Consortium OR JV",
    "Work to be Done Site or Workshop", "Payment terms",
    "Courier Speed Post Submission", "Mandatory Sample Submission",
    "Organization Tender ID", "Scope Classification", "Eligibility Criteria",
    "Checklist",
]

_SUMMARY_ALT = "|".join(
    r"\s+".join(re.escape(w) for w in lab.split())
    for lab in sorted(SUMMARY_LABELS, key=len, reverse=True)
)
SUMMARY_LABEL_RE = re.compile(rf"\b({_SUMMARY_ALT})\b", re.I)

SUMMARY_MARKER_RE = re.compile(
    r"AI\s+Generated\s+Tender\s+Summary|GEM\s+Bid\s+number", re.I)


def _dedupe_repeat(s: str) -> str:
    """Collapse "X X" into "X" - these cells repeat their own value."""
    half = len(s) // 2
    if half > 8 and s[:half].strip() == s[half:].strip():
        return s[:half].strip()
    return s


def _summary_cells(pages: list[Page]) -> tuple[dict[str, str], str]:
    """Return {label_lower: value} from the summary sheet, if there is one.

    Qualification is per FILE, not per page: only the first page carries the
    "AI Generated Tender Summary" banner, while the cost, the EMD and the
    site sit on page 2. Matching per page would silently drop exactly the
    fields the sheet is most useful for.
    """
    by_file: dict[str, list[Page]] = {}
    for pg in _ordered_pages(pages):
        by_file.setdefault(pg.file, []).append(pg)

    cells: dict[str, str] = {}
    ref = ""
    for fname, fpages in by_file.items():
        looks_like = ("tender-summary" in fname.lower()
                      or "tender_summary" in fname.lower()
                      or any(SUMMARY_MARKER_RE.search(p.text) for p in fpages))
        if not looks_like:
            continue
        flat = re.sub(r"\s+", " ", "\n".join(p.text for p in fpages))
        hits = list(SUMMARY_LABEL_RE.finditer(flat))
        if len(hits) < 6:            # not really one of these sheets
            continue
        ref = ref or fpages[0].ref
        for i, m in enumerate(hits):
            end = hits[i + 1].start() if i + 1 < len(hits) else len(flat)
            val = _dedupe_repeat(flat[m.end():end].strip(" :|-–"))
            key = re.sub(r"\s+", " ", m.group(1)).strip().lower()
            if val and key not in cells:
                cells[key] = val
    return cells, ref


def _summary_money(val: str) -> str:
    """Format a summary-sheet money cell, honouring an explicit nil."""
    if EXPLICIT_NIL_RE.search(val):
        m = EXPLICIT_NIL_RE.search(val)
        return m.group(0).strip().title()
    num, disp = parse_amount(val, allow_bare=True)
    return disp or ""


def summary_sheet_extract(pages: list[Page]) -> dict[str, Cand]:
    cells, ref = _summary_cells(pages)
    if not cells:
        return {}

    out: dict[str, Cand] = {}

    def put(field: str, raw: str, conf: str = "high"):
        raw = _clean(raw or "").strip()
        if raw and not _is_junk_line(raw):
            out[field] = Cand(raw, ref, conf)

    for lab in ("estimated cost",):
        if lab in cells:
            put("estimated_cost", _summary_money(cells[lab]))
    if "emd value" in cells:
        put("emd", _summary_money(cells["emd value"]))
    if "bid end date time" in cells:
        put("submission_date", cells["bid end date time"])
    for lab in ("contract period", "completion period"):
        if lab in cells:
            put("period", cells[lab])
            break
    for lab in ("site location", "location"):
        if lab in cells:
            put("location", cells[lab])
            break
    for lab in ("products", "item category"):
        if lab in cells:
            put("purpose", cells[lab])
            break
    if "eligibility criteria" in cells:
        put("eligibility", cells["eligibility criteria"], "medium")
    if "item category" in cells:
        put("scope_of_work", cells["item category"], "medium")
    if "gem bid number" in cells:
        title = cells.get("item category", "")
        tail = title.split(" - ")[-1].strip() if " - " in title else ""
        name = cells["gem bid number"] + (f" - {tail}" if tail else "")
        put("tender_name", name)
    return out


def rules_extract(pages: list[Page], folder_name: str) -> dict[str, Cand]:
    # Read the tender's own documents first. Only if a field is still missing
    # is the standing-conditions template consulted, and what it yields is
    # marked low and labelled, because it describes the department's standard
    # terms rather than this assignment.
    real = [p for p in pages if not p.boilerplate]
    if real and len(real) != len(pages):
        out = _rules_pass(real, folder_name)
        fallback = _rules_pass(pages, folder_name)
        for key, cand in fallback.items():
            if not out.get(key, Cand()).value and cand.value:
                cand.conf = "low"
                cand.value = (cand.value.rstrip()
                              + "\n\n[from the standard contract conditions, "
                                "not this tender - verify]")
                out[key] = cand
        return out
    return _rules_pass(pages, folder_name)


def _rules_pass(pages: list[Page], folder_name: str) -> dict[str, Cand]:
    out: dict[str, Cand] = {}

    for f in MONEY_FIELDS:
        out[f] = _label_lookup(pages, LABELS.get(f, []), "money", f)
    for f in DATE_FIELDS:
        out[f] = _label_lookup(pages, LABELS.get(f, []), "date", f)
    for f in ("period", "location", "purpose", "tender_name"):
        out[f] = _label_lookup(pages, LABELS.get(f, []), "text", f)
    for f, pats in SECTION_HEADINGS.items():
        if f == "eligibility":
            # Criteria span pages; a per-page capture loses most of them.
            out[f] = _capture_section(pages, pats, max_chars=14000,
                                      span_pages=True,
                                      break_re=ELIG_BREAK_RE)
        elif f == "penalty":
            # Break at the very next numbered clause: the SCC lists
            # penalties as one-line table rows, so any run-on swallows
            # the unrelated clauses that follow.
            out[f] = _capture_section(pages, pats, min_before_break=0)
        else:
            out[f] = _capture_section(pages, pats)

    # A stated rate beats a clause heading. "Liquidated Damages / As per
    # GCC" is a pointer; the reader needs the number it points at.
    rate = find_penalty_rate(pages)
    if rate.value:
        head = out.get("penalty", Cand()).value
        if head and head[:60] not in rate.value:
            rate = Cand(rate.value + "\n\n[clause reference: "
                        + _clean(head)[:200] + "]", rate.ref, rate.conf)
        out["penalty"] = rate

    # A GeM bid number is an unambiguous identifier - prefer it.
    for pg in _ordered_pages(pages):
        m = GEM_BID_NO_RE.search(pg.text)
        if m:
            existing = out["tender_name"].value
            extra = f" - {existing}" if existing and m.group(1) not in existing else ""
            out["tender_name"] = Cand(m.group(1) + extra, pg.ref, "high")
            break

    # GeM bids carry no tender document fee. Saying so beats a blank cell that
    # looks like the tool simply failed to find one.
    is_gem = any(GEM_BID_NO_RE.search(p.text) for p in pages[:40])
    if is_gem and not out.get("tender_fees", Cand()).value:
        out["tender_fees"] = Cand(
            "Nil - GeM bids carry no tender document fee",
            "GeM bid document", "medium")

    # ePBG is a percentage of contract value, not a rupee figure - label it as such.
    for pg in _ordered_pages(pages):
        m = EPBG_RE.search(pg.text)
        if m:
            note = f"ePBG {m.group(1)}% of contract value"
            cur = out.get("sd", Cand())
            joined = (cur.value + " | " if cur.value else "") + note
            out["sd"] = Cand(joined, pg.ref, "high")
            break

    # The summary sheet is a structured table, so it outranks anything the
    # regexes scraped out of running prose - except a high-confidence hit,
    # which came off an equally explicit label in the tender itself. For the
    # prose fields the long documents say far more, so only fill a gap there.
    summary = summary_sheet_extract(pages)
    if summary:
        log(f"   read tender summary sheet ({len(summary)} field(s))")
    for f, cand in summary.items():
        cur = out.get(f) or Cand()
        if f in ("eligibility", "scope_of_work", "penalty"):
            if not cur.value:
                out[f] = cand
        elif not cur.value or cur.conf != "high":
            out[f] = cand

    if not out["tender_name"].value:
        out["tender_name"] = Cand(folder_name, "folder name", "low")
    return out


# --------------------------------------------------------------------------
# 4. GEMINI LAYER  -  free-tier LLM pass for judgement fields and gap-filling
# --------------------------------------------------------------------------

PROMPT_HEADER = """You are a highly meticulous Chartered Accountant and Indian Government Tender Analyst.
Below is the complete text of the documents issued for ONE single tender. Page markers
look like <<<FILE: name | PAGE: n>>>.

CRITICAL INSTRUCTIONS: You must read EVERY SINGLE WORD of this document carefully. Remember all details.
Extract the 13 fields below accurately. Rules you MUST follow:
1. STRICT ANTI-HALLUCINATION: Use ONLY what the documents explicitly state. Never guess, infer, or pull numbers from your prior knowledge or generic clauses.
2. If a field is genuinely not stated in the text, you MUST set the value to exactly: NOT FOUND
3. EXHAUSTIVE EXTRACTION: For fields like "Eligibility Criteria" or "Scope of Work", extract the FULL, EXTENSIVE list exactly as written. DO NOT summarize or truncate.
4. Amounts: give the rupee figure in digits (e.g. "Rs. 20,000"). If a percentage is stated, give the percentage and what it is a percentage of.
5. Dates: copy them exactly as written, including time.
6. source_file and page MUST be the exact file name and page number of the marker where you found the value.
7. confidence: "high" when explicitly labeled; "medium" when inferred from context; "low" when unsure.

Field meanings:
1  tender_name      Tender / bid / NIT reference number and its title.
2  location         Place of work, office address, city/district/state.
3  purpose          What is being procured, i.e. the type of audit or service.
4  period           Audit period or contract/engagement duration (FY, months, years).
5  estimated_cost   Tender estimated value / estimated contract cost.
6  assignment_fees  Fee payable for the assignment (audit/professional fee).
7  eligibility      Eligibility / pre-qualification criteria. Extract ALL criteria fully.
8  scope_of_work    Scope of work / terms of reference. Extract ALL points fully.
9  penalty          Penalty, liquidated damages, deduction clauses. Do not mix with banning clauses.
10 emd              Earnest money deposit / bid security.
11 sd               Security deposit / performance guarantee / PBG.
12 tender_fees      Cost of tender document / bid processing or application fee.
13 submission_date  Last date and time for bid submission (the deadline).

Return ONLY a JSON object of this exact shape:
{"fields": {"<field_key>": {"value": "...", "source_file": "...", "page": "...",
"confidence": "high|medium|low"}, ...}}
where field_key is one of: tender_name, location, purpose, period,
estimated_cost, assignment_fees, eligibility, scope_of_work, penalty, emd, sd,
tender_fees, submission_date
"""


# An account can expose dozens of models, most of which cannot read a tender:
# image and speech generators, embedders, and preview builds that answer 503
# far more often than they answer. Every one tried costs a failed request, and
# with backoff between them that is minutes of waiting for nothing.
UNUSABLE_MODEL_RE = re.compile(
    r"embedding|aqa|image|vision|tts|audio|speech|video|omni|live|preview|"
    r"gemma|learnlm|exp-|-exp\b", re.I)

MAX_MODEL_CANDIDATES = 4


def _unusable_model(name: str) -> bool:
    return bool(UNUSABLE_MODEL_RE.search(name))


def _call_groq(api_key: str, models: list[str], prompt: str) -> dict:
    """Call Groq API with automatic model failover (120B -> 27B -> compound)."""
    import requests

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    for model in models:
        try:
            r = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json={
                    "model": model,
                    "messages": [
                        {
                            "role": "system",
                            "content": (
                                "You are an expert chartered accountant and Indian government "
                                "tender analyst. Extract the requested fields strictly in JSON "
                                "format adhering precisely to the JSON schema provided."
                            ),
                        },
                        {"role": "user", "content": prompt},
                    ],
                    "response_format": {"type": "json_object"},
                    "temperature": 0.0,
                },
                timeout=60,
            )
            if r.status_code == 200:
                data = r.json()
                content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
                parsed = _parse_json(content)
                if parsed:
                    log(f"     Groq AI succeeded on {model}")
                    return parsed
            elif r.status_code == 429:
                log(f"   - Groq model {model} rate limited, trying next model")
                continue
            else:
                log(f"   ! Groq {model} returned status {r.status_code}: {r.text[:80]}")
        except Exception as e:
            log(f"   ! Groq {model} exception: {str(e)[:80]}")
            continue
    return {}


def _gemini_client():
    """Build a Gemini client and pick an active model this key can use."""
    try:
        from google import genai
    except Exception as e:
        log(f"   ! Gemini SDK not installed ({e})")
        return None, None

    key = (CONFIG.get("gemini_api_key") or os.environ.get("GEMINI_API_KEY") or "").strip()
    if not key:
        return None, None

    try:
        client = genai.Client(api_key=key)
        
        # Dynamically fetch available models to avoid trying unavailable ones
        try:
            available = [
                m.name.replace("models/", "") 
                for m in client.models.list() 
                if "flash" in m.name.lower() and "vision" not in m.name.lower()
            ]
            candidates = sorted(available, reverse=True)
            if not candidates:
                raise ValueError("No flash models found")
        except Exception:
            candidates = list(CONFIG.get("gemini_models") or [
                "gemini-3.6-flash", "gemini-3.7-flash", "gemini-3.5-flash-lite", "gemini-flash-latest"
            ])
            
        for model in candidates:
            try:
                client.models.generate_content(model=model, contents="ping")
                log(f"   Gemini AI ready: {model}")
                return client, [model] + [m for m in candidates if m != model]
            except Exception as e:
                msg = str(e)
                if _overloaded(msg):
                    log(f"   - {model} is busy; keeping as fallback")
                    continue
                if "401" in msg or "API_KEY_INVALID" in msg or "PERMISSION_DENIED" in msg:
                    log(f"   ! Gemini key rejected: {msg[:70]}")
                    return None, None
                # Silently skip unavailable models without logging to avoid console spam
                pass
    except Exception as e:
        log(f"   ! Gemini client init error: {e}")

    return None, None


def _corpus(pages: list[Page]) -> str:
    # Prioritize real tender documents over boilerplate general conditions
    real_pages = [p for p in pages if not p.boilerplate]
    boiler_pages = [p for p in pages if p.boilerplate]
    sorted_pages = _ordered_pages(real_pages) + _ordered_pages(boiler_pages)

    parts = []
    for pg in sorted_pages:
        parts.append(f"<<<FILE: {pg.file} | PAGE: {pg.page}"
                     f"{' | OCR' if pg.ocr else ''}>>>\n{pg.text}")
    return "\n\n".join(parts)


def _chunks(text: str, chunk_size: int = 35_000, max_chunks: int = 8) -> list[str]:
    if len(text) <= chunk_size:
        return [text]
    out, i = [], 0
    while i < len(text) and len(out) < max_chunks:
        cut = text.rfind("\n<<<FILE:", i + int(chunk_size * 0.6), i + chunk_size)
        if cut == -1:
            cut = min(i + chunk_size, len(text))
        out.append(text[i:cut])
        i = cut
    return out


def _parse_json(txt: str) -> dict:
    txt = (txt or "").strip()
    txt = re.sub(r"^```(?:json)?|```$", "", txt, flags=re.M).strip()
    try:
        return json.loads(txt)
    except Exception:
        pass
    s, e = txt.find("{"), txt.rfind("}")
    if s != -1 and e > s:
        try:
            return json.loads(txt[s:e + 1])
        except Exception:
            pass
    return {}


def _overloaded(msg: str) -> bool:
    """503 / UNAVAILABLE - the model is busy. A different model may be free."""
    m = msg.lower()
    return ("503" in m or "unavailable" in m or "overloaded" in m
            or "high demand" in m)


def _rate_limited(msg: str) -> bool:
    """429 - your free-tier quota. Every model shares it, so waiting is the fix."""
    m = msg.lower()
    return "429" in m or "resource_exhausted" in m or "quota" in m


def _dead_model(msg: str) -> bool:
    m = msg.lower()
    return "404" in m or "not found" in m or "not supported" in m


def _call_gemini(client, models: list[str], prompt: str) -> dict:
    """Try each model in turn; rotate on overload, wait only on real quota limits."""
    if isinstance(models, str):
        models = [models]
    live = [m for m in models if m]
    delay = 5
    for rnd in range(CONFIG["gemini_retries"]):
        if not live:
            break
        quota_hit = False
        for model in list(live):
            try:
                resp = client.models.generate_content(
                    model=model, contents=prompt,
                    config={"response_mime_type": "application/json",
                            "temperature": 0},
                )
                out = _parse_json(getattr(resp, "text", ""))
                if out:
                    if rnd or model != models[0]:
                        log(f"     Gemini succeeded on {model}")
                    return out
                log(f"   ! {model} returned nothing parseable")
            except Exception as e:
                msg = str(e)
                if _overloaded(msg):
                    log(f"   - {model} overloaded (503), trying next model")
                    continue
                if _dead_model(msg):
                    log(f"   - {model} not available on this key, dropping it")
                    live.remove(model)
                    if model in models:
                        models.remove(model)
                    continue
                if _rate_limited(msg):
                    log(f"   - free-tier quota hit on {model}; pausing before retry")
                    quota_hit = True
                    break
                log(f"   ! {model} failed: {msg[:120]}")
                continue
        if rnd < CONFIG["gemini_retries"] - 1:
            why = "quota" if quota_hit else "all models busy"
            log(f"     {why}; waiting {delay}s before round {rnd + 2}")
            time.sleep(delay)
            delay = min(delay * 2, 30)
    return {}


def _merge_ai(results: list[dict]) -> dict[str, Cand]:
    """Merge per-chunk AI results: longest wins for prose, best confidence for scalars."""
    rank = {"high": 3, "medium": 2, "low": 1, "": 0}
    merged: dict[str, Cand] = {}
    for res in results:
        fields = (res or {}).get("fields") or {}
        for key, _ in FIELDS:
            item = fields.get(key)
            if not isinstance(item, dict):
                continue
            val = str(item.get("value") or "").strip()
            if not val or val.upper().startswith("NOT FOUND"):
                continue
            ref = str(item.get("source_file") or "").strip()
            pageno = str(item.get("page") or "").strip()
            if ref and pageno:
                ref = f"{ref} (p.{pageno})"
            conf = str(item.get("confidence") or "medium").lower()
            cand = Cand(val, ref or "AI", conf if conf in rank else "medium")
            cur = merged.get(key)
            if cur is None:
                merged[key] = cand
            elif key in PROSE_FIELDS:
                if len(val) > len(cur.value):
                    merged[key] = cand
            elif rank[cand.conf] > rank[cur.conf]:
                merged[key] = cand
    return merged


def ai_extract(pages: list[Page], gemini_client, gemini_models, groq_key: str = "", groq_models: list[str] = None):
    """Unified AI extraction with Groq 120B and Gemini Flash."""
    groq_key = (groq_key or CONFIG.get("groq_api_key") or os.environ.get("GROQ_API_KEY") or "").strip()
    groq_models = groq_models or list(CONFIG.get("groq_models") or [
        "openai/gpt-oss-120b", "qwen/qwen3.8-27b", "groq/compound", "openai/gpt-oss-20b"
    ])

    use_groq = bool(groq_key) and CONFIG.get("use_groq", True)
    use_gemini = bool(gemini_client) and CONFIG.get("use_gemini", True)

    if not use_groq and not use_gemini:
        if CONFIG.get("require_ai", True):
            log("   ! Compulsory AI is enabled, but no valid Groq or Gemini API key was found!")
        return {}, False

    provider = CONFIG.get("ai_provider", "auto").lower()
    chunk_size = 15_000 if (use_groq and provider in ("auto", "groq")) else 120_000
    corpus = _corpus(pages)
    chunks = _chunks(corpus, chunk_size=chunk_size, max_chunks=CONFIG.get("gemini_max_chunks", 8))
    results = []

    for n, ch in enumerate(chunks, start=1):
        prompt = PROMPT_HEADER + "\n\n===== DOCUMENTS =====\n" + ch
        res = {}

        # 1. Try Groq (Ultra-fast 120B / 27B)
        if use_groq and provider in ("auto", "groq", "both"):
            log(f"   Groq AI pass {n}/{len(chunks)} ({len(ch):,} chars)")
            res = _call_groq(groq_key, groq_models, prompt)

        # 2. Fall back to Gemini if Groq failed or provider is gemini/both
        if not res and use_gemini and provider in ("auto", "gemini", "both"):
            log(f"   Gemini AI pass {n}/{len(chunks)} ({len(ch):,} chars)")
            res = _call_gemini(gemini_client, gemini_models, prompt)

        if res:
            results.append(res)
        else:
            log(f"   ! AI pass {n} could not be completed")

        if n < len(chunks):
            missing = [k for k, _ in FIELDS if k not in _merge_ai(results)]
            if not missing:
                log(f"     all 13 fields found; skipping {len(chunks) - n} remaining pass(es)")
                break

    ai_ok = any(bool(r) for r in results)
    return _merge_ai(results), ai_ok


# --------------------------------------------------------------------------
# 5. MERGE  -  combine rules + AI, flag disagreements, never guess
# --------------------------------------------------------------------------

@dataclass
class Result:
    value: str = NOT_FOUND
    ref: str = ""
    conf: str = ""
    rules_value: str = ""
    ai_value: str = ""
    flag: str = ""


def _num(text: str) -> float | None:
    m = re.search(r"(\d[\d,]*(?:\.\d+)?)", text or "")
    if not m:
        return None
    try:
        val = float(m.group(1).replace(",", ""))
    except ValueError:
        return None
    # Only a multiplier immediately after the digits scales them. Looking
    # anywhere in the string wrongly scaled "Rs. 2,50,000/- (Rupees Two Lakhs
    # fifty thousand only)" by another 100,000.
    tail = (text or "")[m.end(): m.end() + 14].lower()
    if re.match(r"\s*(?:/-)?\s*(?:crores?|cr\b)", tail):
        val *= 1e7
    elif re.match(r"\s*(?:/-)?\s*(?:lakh?s?|lacs?)", tail):
        val *= 1e5
    return val


def merge(rules: dict[str, Cand], ai: dict[str, Cand],
          ai_ran: bool = True) -> dict[str, Result]:
    out: dict[str, Result] = {}
    for key, _ in FIELDS:
        r = rules.get(key) or Cand()
        a = ai.get(key) or Cand()
        res = Result(rules_value=r.value, ai_value=a.value)

        if a.value:
            res.value, res.ref, res.conf = a.value, (a.ref or r.ref), a.conf
        elif r.value:
            res.value, res.ref, res.conf = r.value, r.ref, r.conf
        else:
            res.value, res.conf = NOT_FOUND, ""
            res.flag = "NOT FOUND"
            out[key] = res
            continue

        # Cross-check the money and date fields: two independent readers.
        if key in MONEY_FIELDS and r.value and a.value:
            rn, an = _num(r.value), _num(a.value)
            if rn and an and abs(rn - an) > max(1.0, 0.01 * max(rn, an)):
                res.flag = "CHECK - readers disagree"
                res.value += f"   [rules read: {r.value.split('  [')[0]}]"
                res.conf = "low"
        if key in DATE_FIELDS and r.value and a.value:
            rd = re.sub(r"\D", "", r.value)[:8]
            ad = re.sub(r"\D", "", a.value)[:8]
            if rd and ad and rd != ad:
                res.flag = "CHECK - readers disagree"
                res.value += f"   [rules read: {r.value}]"
                res.conf = "low"

        # One reader is not a cross-check. Reporting "high" confidence on a
        # money field that only the rules layer saw is what let a wrong
        # security deposit through unflagged in tender 2026-9533.
        if not ai_ran and key in (MONEY_FIELDS | DATE_FIELDS):
            if res.conf == "high":
                res.conf = "medium"
            if not res.flag:
                res.flag = "single reader - AI cross-check did not run"

        if not res.flag and res.conf == "low":
            res.flag = "low confidence"
        if not res.flag and any(w in (res.ref or "").upper() for w in ("OCR",)):
            res.flag = "from OCR - verify digits"
        out[key] = res
    return out


# --------------------------------------------------------------------------
# 6. EXCEL OUTPUT
# --------------------------------------------------------------------------

LEGEND = [
    ("How to read this workbook", ""),
    ("Tender Summary", "One row per tender folder, the 13 requested fields."),
    ("Evidence", "For every field: the source file and page it came from, plus "
                 "what each of the two independent readers (rules and AI) read. "
                 "Use this to verify a value in seconds."),
    ("Documents Read", "Which files were opened, how many pages, how many "
                       "needed OCR."),
    ("", ""),
    ("Cell colours", ""),
    ("Amber", "Needs your eye: the two readers disagreed, confidence was low, "
              "or the value came off a scanned page via OCR."),
    ("Red", "Not stated anywhere in the documents that were read."),
    ("", ""),
    ("Important", "This tool eliminates the typing, not the review. Always "
                  "confirm EMD, fees, and the submission deadline against the "
                  "source page before acting on them."),
    ("Single reader", "Without a Gemini API key only the rules layer runs. "
                      "Money and date fields are then capped at medium "
                      "confidence and flagged, because nothing "
                      "cross-checked them."),
    ("Percentages", "Security deposit and performance guarantees are "
                    "normally a percentage of contract value, so that is "
                    "what is reported. A rupee figure sitting next to a "
                    "phrase like \"up to\" is a threshold, not the amount "
                    "payable, and is flagged as such."),
    ("Excluded", "WORKING FOLDER is skipped by design - it holds your own "
                 "draft submissions, not the department's tender documents."),
]


def write_excel(rows: list[dict], out_path: Path) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter

    hdr_fill = PatternFill("solid", fgColor="1F3864")
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    amber = PatternFill("solid", fgColor="FFF2CC")
    red = PatternFill("solid", fgColor="FCE4E4")
    thin = Side(style="thin", color="BFBFBF")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    wb = Workbook()

    # ---- Sheet 1: Tender Summary -----------------------------------------
    ws = wb.active
    ws.title = "Tender Summary"
    headers = ["Tender Folder"] + [label for _, label in FIELDS] + ["Needs Review"]
    ws.append(headers)
    for c in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=c)
        cell.fill, cell.font, cell.border = hdr_fill, hdr_font, box
        cell.alignment = Alignment(wrap_text=True, vertical="center",
                                   horizontal="center")
    ws.row_dimensions[1].height = 34

    for r_i, row in enumerate(rows, start=2):
        res: dict[str, Result] = row["results"]
        flags = [label for key, label in FIELDS if res[key].flag]
        ws.cell(row=r_i, column=1, value=row["tender"])
        for c_i, (key, _) in enumerate(FIELDS, start=2):
            val = res[key].value
            cell = ws.cell(row=r_i, column=c_i,
                           value=val[:2000] if val else NOT_FOUND)
            if res[key].flag == "NOT FOUND":
                cell.fill = red
            elif res[key].flag:
                cell.fill = amber
        ws.cell(row=r_i, column=len(headers),
                value=(f"{len(flags)} field(s): " +
                       ", ".join(f.split(". ", 1)[-1] for f in flags))
                if flags else "clean")
        for c in range(1, len(headers) + 1):
            ws.cell(row=r_i, column=c).alignment = wrap
            ws.cell(row=r_i, column=c).border = box
        ws.row_dimensions[r_i].height = 150

    widths = [18, 30, 26, 24, 16, 20, 20, 46, 52, 34, 20, 20, 18, 24, 30]
    for i, w in enumerate(widths[:len(headers)], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "B2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{max(2, len(rows) + 1)}"

    # ---- Sheet 2: Evidence -----------------------------------------------
    ev = wb.create_sheet("Evidence")
    ev_head = ["Tender", "Field", "Final Value", "Source (file, page)",
               "Confidence", "Rules layer read", "AI layer read", "Flag"]
    ev.append(ev_head)
    for c in range(1, len(ev_head) + 1):
        cell = ev.cell(row=1, column=c)
        cell.fill, cell.font, cell.border = hdr_fill, hdr_font, box
    r = 2
    for row in rows:
        res = row["results"]
        for key, label in FIELDS:
            x = res[key]
            for c_i, v in enumerate([row["tender"], label, x.value, x.ref,
                                     x.conf, x.rules_value, x.ai_value,
                                     x.flag], start=1):
                cell = ev.cell(row=r, column=c_i, value=str(v)[:1500])
                cell.alignment = wrap
                cell.border = box
                if x.flag == "NOT FOUND":
                    cell.fill = red
                elif x.flag:
                    cell.fill = amber
            ev.row_dimensions[r].height = 42
            r += 1
    for i, w in enumerate([18, 26, 60, 34, 12, 40, 40, 24], start=1):
        ev.column_dimensions[get_column_letter(i)].width = w
    ev.freeze_panes = "C2"

    # ---- Sheet 3: Documents Read -----------------------------------------
    dr = wb.create_sheet("Documents Read")
    dr.append(["Tender", "File", "Pages read", "Pages needing OCR", "Note"])
    for c in range(1, 6):
        cell = dr.cell(row=1, column=c)
        cell.fill, cell.font, cell.border = hdr_fill, hdr_font, box
    r = 2
    for row in rows:
        for f in row["files"]:
            for c_i, v in enumerate([row["tender"], f["name"], f["pages"],
                                     f["ocr"], f["note"]], start=1):
                dr.cell(row=r, column=c_i, value=v).border = box
            r += 1
    for i, w in enumerate([18, 62, 12, 20, 40], start=1):
        dr.column_dimensions[get_column_letter(i)].width = w
    dr.freeze_panes = "A2"

    # ---- Sheet 4: Read Me ------------------------------------------------
    rm = wb.create_sheet("Read Me")
    for a, b in LEGEND:
        rm.append([a, b])
    for i, row in enumerate(rm.iter_rows(min_row=1, max_col=2), start=1):
        row[0].font = Font(bold=True)
        row[1].alignment = Alignment(wrap_text=True, vertical="top")
    rm.column_dimensions["A"].width = 28
    rm.column_dimensions["B"].width = 92

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(out_path))
    return out_path


# --------------------------------------------------------------------------
# 7. DRIVER
# --------------------------------------------------------------------------

def _gather_files(folder: Path) -> list[Path]:
    files = []
    for p in sorted(folder.rglob("*")):
        if p.is_dir():
            continue
        if any(_excluded_dir(part) for part in p.relative_to(folder).parts[:-1]):
            continue
        if _excluded_file(p.name) or p.name.startswith("_conv"):
            continue
        if p.suffix.lower() in READERS:
            files.append(p)
    return sorted(files, key=lambda p: (-file_priority(p.name), p.name))


def discover_tenders(root: Path) -> list[tuple[str, list[Path]]]:
    """Each top-level sub-folder of root is one tender."""
    tenders = []
    subdirs = [d for d in sorted(root.iterdir())
               if d.is_dir() and not _excluded_dir(d.name)]
    for d in subdirs:
        files = _gather_files(d)
        if files:
            tenders.append((d.name, files))
        else:
            log(f"   - {d.name}: no readable documents found")
    loose = [p for p in sorted(root.glob("*")) if p.is_file()
             and p.suffix.lower() in READERS and not _excluded_file(p.name)]
    # Documents sitting directly in the root used to be read only when there
    # were no sub-folders at all, so a folder holding both tender sub-folders
    # and a few loose files silently dropped the loose ones - no warning, they
    # simply never appeared in the summary.
    if loose:
        tenders.append((root.name, loose))
    return tenders


def _drive_id_for(path: Path) -> str:
    """Drive id of a downloaded file, matched on its full relative path.

    Matching on the bare file name would collide: several tenders ship a
    'GeM _ Upload Documents.pdf', and the first id found would then stand in
    for all of them, so a change to one would not invalidate the others.
    """
    p = str(path).replace("\\", "/")
    for rel, fid in DRIVE_IDS.items():
        r = rel.replace("\\", "/")
        if p == r or p.endswith("/" + r):
            return fid
    return ""


def _tender_hash(files: list[Path]) -> str:
    """Fingerprint a tender's inputs.

    Size alone missed real edits - a corrected page that happens to occupy
    the same number of bytes, and every Drive-side replacement, since the
    local copy was never refreshed. Modification time and the Drive file id
    are folded in so a document that has actually changed moves the key.
    """
    h = hashlib.sha1()
    for f in sorted(files, key=lambda p: str(p).lower()):
        try:
            st = f.stat()
            h.update(f"{f.name}:{st.st_size}:{int(st.st_mtime)}".encode())
        except OSError:
            h.update(f.name.encode())
        fid = _drive_id_for(f)
        if fid:
            h.update(fid.encode())
    return h.hexdigest()[:16]


def process_tender(name: str, files: list[Path], client, models, groq_key: str = "", groq_models: list[str] = None) -> dict:
    log(f"\n== {name}  ({len(files)} document(s))")
    pages: list[Page] = []
    file_log = []
    for f in files:
        log(f"   reading {f.name}")
        pgs = read_any(f)
        pages.extend(pgs)
        ocr_n = sum(1 for p in pgs if p.ocr)
        file_log.append({
            "name": f.name, "pages": len(pgs), "ocr": ocr_n,
            "note": ("no text extracted - likely an image-only scan and OCR "
                     "could not read it" if not pgs else
                     ("scanned, read via OCR" if ocr_n else "")),
        })
    if not pages:
        log("   ! no readable text in this tender folder")
        empty = {k: Result(NOT_FOUND, "", "", "", "", "NOT FOUND")
                 for k, _ in FIELDS}
        return {"tender": name, "results": empty, "files": file_log,
                "ai_ok": True}

    total_chars = sum(len(p.text) for p in pages)
    log(f"   {len(pages)} page(s), {total_chars:,} characters of text")
    mark_boilerplate(pages)
    rules = rules_extract(pages, name)
    ai, ai_ok = ai_extract(pages, client, models, groq_key, groq_models)
    
    ai_active = bool(client or groq_key)
    if CONFIG.get("require_ai", True) and ai_active and not ai_ok:
        log("   ! Compulsory AI extraction failed; falling back to rules")

    results = merge(rules, ai, ai_ran=ai_active and ai_ok)
    found = sum(1 for k, _ in FIELDS if results[k].value != NOT_FOUND)
    
    provider_labels = []
    if groq_key and CONFIG.get("use_groq", True):
        provider_labels.append("Groq 120B")
    if client and CONFIG.get("use_gemini", True):
        provider_labels.append("Gemini Flash")
    
    lbl = " + ".join(provider_labels) if provider_labels else "Rules only"
    if not ai_ok and ai_active:
        lbl += " (AI unreachable; rules fallback)"
    log(f"   -> {found}/13 fields found [{lbl}]")
    return {"tender": name, "results": results, "files": file_log,
            "ai_ok": ai_ok}


# Stale caches from an older, weaker engine must never be silently reused.
# This used to be a hand-bumped constant, which is a trap: change the
# extraction logic, forget the bump, and every tender replays yesterday's
# wrong answers in six seconds while looking like a successful run. The
# engine's own source is hashed instead, so any edit to the extraction logic
# invalidates the cache automatically.
CACHE_VERSION = "v5"


def _engine_fingerprint() -> str:
    try:
        src = Path(__file__).read_bytes()
        return CACHE_VERSION + "-" + hashlib.md5(src).hexdigest()[:8]
    except Exception:
        # Embedded/exec'd copies have no __file__; fall back to the constant.
        return CACHE_VERSION


ENGINE_FP = _engine_fingerprint()


def _cache_path(work: Path, name: str, h: str) -> Path:
    safe = re.sub(r"[^\w.-]", "_", name)
    return work / "cache" / f"{safe}__{ENGINE_FP}_{h}.json"


def _to_json(row: dict) -> str:
    return json.dumps({
        "tender": row["tender"], "files": row["files"],
        "ai_ok": row.get("ai_ok", True),
        "results": {k: vars(v) for k, v in row["results"].items()},
    }, ensure_ascii=False)


def _from_json(txt: str) -> dict:
    d = json.loads(txt)
    return {"tender": d["tender"], "files": d["files"],
            "ai_ok": d.get("ai_ok", True),
            "results": {k: Result(**v) for k, v in d["results"].items()}}


# Populated by run() so a caller (e.g. the notebook, to render the dashboard
# inline) can get at the per-tender data without re-parsing the Excel.
LAST_ROWS: list = []


def run() -> Path:
    global LAST_ROWS
    t0 = time.time()
    work = Path(CONFIG["work_dir"])
    (work / "cache").mkdir(parents=True, exist_ok=True)

    print("=" * 68)
    print("TENDER SUMMARY EXTRACTOR")
    print("=" * 68)

    if CONFIG["source_mode"] == "drive_link":
        print("\n[1/4] Fetching documents from Google Drive ...")
        root = download_drive_folder(CONFIG["drive_folder_url"], work / "docs")
    else:
        root = Path(CONFIG["local_folder"]).expanduser()
        print(f"\n[1/4] Using local folder: {root}")
        if not root.exists():
            raise FileNotFoundError(f"Folder not found: {root}")

    print("\n[2/4] Finding tender folders ...")
    tenders = discover_tenders(root)
    if not tenders:
        raise RuntimeError("No tender folders with readable documents found.")
    print(f"   {len(tenders)} tender(s): " + ", ".join(t[0] for t in tenders))

    print("\n[3/4] Reading documents and extracting fields ...")
    groq_key = (CONFIG.get("groq_api_key") or os.environ.get("GROQ_API_KEY") or "").strip()
    groq_models = CONFIG.get("groq_models") or [
        "openai/gpt-oss-120b", "qwen/qwen3.8-27b", "groq/compound", "openai/gpt-oss-20b"
    ]
    client, models = (_gemini_client() if CONFIG.get("use_gemini", True) else (None, None))

    if groq_key:
        print("   * Groq AI active (120B reasoning engine)")
    if client:
        print(f"   * Gemini AI active (model: {models[0] if models else 'flash'})")
    if not groq_key and not client:
        if CONFIG.get("require_ai", True):
            print("   ! WARNING: AI is set to compulsory, but no API keys were configured!")
        else:
            print("   - Running in rules-only mode (offline).")

    rows = []
    for name, files in tenders:
        cache = _cache_path(work, name, _tender_hash(files))
        if CONFIG["use_cache"] and cache.exists():
            try:
                cached = _from_json(cache.read_text(encoding="utf-8"))
                # Only reuse a cached tender that was fully extracted. One
                # whose AI pass failed gets retried instead of being frozen in.
                if cached.get("ai_ok", True):
                    rows.append(cached)
                    log(f"\n== {name}: unchanged since last run, using cache")
                    continue
                log(f"\n== {name}: cached run had no AI - redoing it")
            except Exception:
                pass
        try:
            row = process_tender(name, files, client, models, groq_key, groq_models)
        except Exception:
            log(f"   ! {name} failed:\n{traceback.format_exc(limit=3)}")
            row = {"tender": name, "files": [], "ai_ok": True,
                   "results": {k: Result(NOT_FOUND, "", "", "", "",
                                         "processing error") for k, _ in FIELDS}}
        rows.append(row)
        if row.get("ai_ok", True):
            try:
                cache.write_text(_to_json(row), encoding="utf-8")
            except Exception:
                pass

    LAST_ROWS = rows
    print("\n[4/4] Writing Excel and dashboard ...")
    out = write_excel(rows, Path(CONFIG["output_xlsx"]))
    stamp = "Run " + time.strftime("%d %b %Y, %H:%M")
    dash = write_dashboard(rows, Path(CONFIG["output_html"]), stamp=stamp)
    print(f"   dashboard: {dash}  (double-click to open in any browser)")

    total = len(rows) * len(FIELDS)
    found = sum(1 for r in rows for k, _ in FIELDS
                if r["results"][k].value != NOT_FOUND)
    flagged = sum(1 for r in rows for k, _ in FIELDS if r["results"][k].flag)
    print("\n" + "=" * 68)
    print(f"DONE in {time.time() - t0:.0f}s   ->  {out}")
    print(f"   tenders processed : {len(rows)}")
    print(f"   fields filled     : {found}/{total}")
    print(f"   fields to review  : {flagged}  (amber/red cells)")
    no_ai = [r["tender"] for r in rows if not r.get("ai_ok", True)]
    if no_ai:
        print(f"   AI unreachable for: {', '.join(no_ai)}")
        print("   -> these used the rules layer only. They were not cached, so")
        print("      simply run this again later and only they get reprocessed.")
    print("=" * 68)
    return out


# --------------------------------------------------------------------------
# 8. HTML DASHBOARD  -  double-click to open, no software, no coding
# --------------------------------------------------------------------------

DASH_STYLE = """
<style>
  @import url('https://fonts.googleapis.com/css2?family=Source+Serif+4:opsz,wght@8..60,400;8..60,600;8..60,700&family=IBM+Plex+Sans:wght@400;500;600&family=IBM+Plex+Mono:wght@400;500;600&display=swap');

  :root {
    --ground:      #F1F4F3;
    --surface:     #FFFFFF;
    --surface-sub: #E9EEEC;
    --line:        #D3DCD9;
    --line-soft:   #E4EAE8;
    --ink:         #10262C;
    --ink-soft:    #46605E;
    --ink-faint:   #7B918D;
    --accent:      #14655A;
    --accent-soft: #DDEAE6;
    --ok:          #2C7150;
    --warn:        #9A6A11;
    --warn-bg:     #FBF1DC;
    --crit:        #9E3527;
    --crit-bg:     #F9E5E1;
    --shadow:      0 1px 2px rgba(16,38,44,.06), 0 6px 18px rgba(16,38,44,.05);
    --shadow-hover: 0 2px 4px rgba(16,38,44,.09), 0 16px 30px rgba(16,38,44,.11);
  }
  :root:not([data-theme="light"]) { }
  @media (prefers-color-scheme: dark) {
    :root:not([data-theme="light"]) {
      --ground:      #0C1A1E;
      --surface:     #132A2F;
      --surface-sub: #1B363B;
      --line:        #27474C;
      --line-soft:   #1F3B40;
      --ink:         #E7EEEC;
      --ink-soft:    #A9C0BC;
      --ink-faint:   #7B948F;
      --accent:      #5BB9A6;
      --accent-soft: #17403C;
      --ok:          #63BA8D;
      --warn:        #D9A94A;
      --warn-bg:     #33290F;
      --crit:        #E2867A;
      --crit-bg:     #3A1D18;
      --shadow:      0 1px 2px rgba(0,0,0,.3), 0 8px 24px rgba(0,0,0,.28);
      --shadow-hover: 0 2px 6px rgba(0,0,0,.38), 0 18px 34px rgba(0,0,0,.34);
    }
  }
  :root[data-theme="dark"] {
    --ground:      #0C1A1E;
    --surface:     #132A2F;
    --surface-sub: #1B363B;
    --line:        #27474C;
    --line-soft:   #1F3B40;
    --ink:         #E7EEEC;
    --ink-soft:    #A9C0BC;
    --ink-faint:   #7B948F;
    --accent:      #5BB9A6;
    --accent-soft: #17403C;
    --ok:          #63BA8D;
    --warn:        #D9A94A;
    --warn-bg:     #33290F;
    --crit:        #E2867A;
    --crit-bg:     #3A1D18;
    --shadow:      0 1px 2px rgba(0,0,0,.3), 0 8px 24px rgba(0,0,0,.28);
    --shadow-hover: 0 2px 6px rgba(0,0,0,.38), 0 18px 34px rgba(0,0,0,.34);
  }

  * { box-sizing: border-box; }
  html { background: var(--ground); }
  body {
    margin: 0; background: var(--ground); color: var(--ink);
    font-family: 'IBM Plex Sans', system-ui, -apple-system, sans-serif;
    font-size: 15px; line-height: 1.55;
    -webkit-font-smoothing: antialiased;
    border-top: 3px solid var(--accent);
  }
  .wrap { max-width: 1240px; margin: 0 auto; padding: 30px 24px 72px; }

  /* ---------- masthead ---------- */
  .mast { display: flex; flex-wrap: wrap; align-items: flex-end;
          justify-content: space-between; gap: 16px;
          padding-bottom: 16px; border-bottom: 1px solid var(--line); position: relative; }
  .mast::after { content: ""; position: absolute; left: 0; right: 0; bottom: -1px;
                 height: 2px; width: 64px; background: var(--accent); }
  .brand { display: flex; align-items: center; gap: 14px; }
  .mark { flex: none; display: block; border-radius: 9px; }
  h1 { font-family: 'Source Serif 4', Georgia, serif; font-weight: 700;
       font-size: clamp(28px, 4vw, 40px); line-height: 1.1; margin: 0;
       letter-spacing: -.01em; text-wrap: balance; }
  .mast .sub { color: var(--ink-soft); font-size: 13.5px; margin-top: 5px; }
  .runstamp { font-family: 'IBM Plex Mono', monospace; font-size: 11.5px;
              color: var(--ink-soft); text-align: right; white-space: nowrap;
              padding: 5px 11px; border-radius: 999px; background: var(--surface-sub);
              border: 1px solid var(--line-soft); }

  /* ---------- summary tiles ---------- */
  .tiles { display: grid; gap: 12px; margin: 22px 0 8px;
           grid-template-columns: repeat(auto-fit, minmax(210px, 1fr)); }
  .tile { display: flex; gap: 12px; align-items: flex-start;
          background: var(--surface); border: 1px solid var(--line-soft);
          border-radius: 6px; padding: 14px 16px; box-shadow: var(--shadow); }
  .tile .icon { flex: none; width: 30px; height: 30px; border-radius: 7px;
                display: flex; align-items: center; justify-content: center;
                background: var(--accent-soft); color: var(--accent); }
  .tile .icon svg { width: 17px; height: 17px; }
  .tile .k { font-size: 10.5px; letter-spacing: .09em; text-transform: uppercase;
             color: var(--ink-faint); font-weight: 600; }
  .tile .v { font-family: 'IBM Plex Mono', monospace; font-variant-numeric: tabular-nums;
             font-size: 25px; font-weight: 600; margin-top: 5px; line-height: 1.1; }
  .tile .n { font-size: 12px; color: var(--ink-soft); margin-top: 3px; }
  .tile.alert { border-color: var(--warn); background: var(--warn-bg); }
  .tile.alert .v { color: var(--warn); }
  .tile.alert .icon { background: rgba(0,0,0,.06); color: var(--warn); }

  /* ---------- controls ---------- */
  .controls { display: flex; flex-wrap: wrap; gap: 10px; align-items: center;
              margin: 22px 0 18px; padding: 12px; background: var(--surface-sub);
              border: 1px solid var(--line-soft); border-radius: 4px; }
  input[type=search], select {
    font: inherit; font-size: 14px; color: var(--ink); background: var(--surface);
    border: 1px solid var(--line); border-radius: 3px; padding: 8px 10px; }
  input[type=search] { flex: 1 1 240px; min-width: 180px; }
  .chk { display: inline-flex; align-items: center; gap: 7px; font-size: 13.5px;
         color: var(--ink-soft); cursor: pointer; user-select: none; }
  .chk input { accent-color: var(--accent); width: 15px; height: 15px; }
  :is(input, select, summary, button):focus-visible {
    outline: 2px solid var(--accent); outline-offset: 2px; }
  .count { margin-left: auto; font-size: 12.5px; color: var(--ink-faint);
           font-family: 'IBM Plex Mono', monospace; }

  /* ---------- tender cards ---------- */
  .cards { display: grid; gap: 16px;
           grid-template-columns: repeat(auto-fit, minmax(500px, 1fr)); }
  .card { background: var(--surface); border: 1px solid var(--line-soft);
          border-left: 4px solid var(--ink-faint); border-radius: 6px;
          box-shadow: var(--shadow); overflow: hidden; display: flex;
          flex-direction: column; min-width: 0;
          animation: cardIn .28s ease both;
          transition: transform .15s ease, box-shadow .15s ease; }
  .card:hover { transform: translateY(-2px); box-shadow: var(--shadow-hover); }
  @keyframes cardIn { from { opacity: 0; transform: translateY(7px); }
                      to { opacity: 1; transform: none; } }
  .card.urgent { border-left-color: var(--crit); }
  .card.soon   { border-left-color: var(--warn); }
  .card.open   { border-left-color: var(--ok); }
  .card.past   { border-left-color: var(--ink-faint); opacity: .72; }

  .chead { display: flex; flex-wrap: wrap; gap: 12px; align-items: flex-start;
           padding: 16px 18px 12px; border-bottom: 1px solid var(--line-soft); }
  .chead .id { font-family: 'IBM Plex Mono', monospace; font-size: 11.5px;
               color: var(--ink-faint); letter-spacing: .04em; }
  .chead h2 { font-family: 'Source Serif 4', Georgia, serif; font-size: 19px;
              font-weight: 600; margin: 3px 0 0; line-height: 1.28;
              text-wrap: balance; }
  .chead .grow { flex: 1 1 320px; min-width: 0; }
  .chips { display: flex; flex-wrap: wrap; gap: 6px; margin-top: 8px; }
  .chip { font-size: 11.5px; padding: 3px 9px; border-radius: 999px;
          background: var(--accent-soft); color: var(--accent); font-weight: 600;
          border: 1px solid transparent; }
  .chip.plain { background: transparent; border-color: var(--line);
                color: var(--ink-soft); font-weight: 500; }
  .due { text-align: right; flex: 0 0 auto; }
  .due .d { font-family: 'IBM Plex Mono', monospace; font-size: 14.5px;
            font-weight: 600; font-variant-numeric: tabular-nums; }
  .pill { display: inline-block; margin-top: 5px; font-size: 11.5px;
          font-weight: 600; padding: 3px 10px; border-radius: 999px;
          border: 1px solid currentColor; }
  .pill.urgent { color: var(--crit); background: var(--crit-bg); }
  .pill.soon   { color: var(--warn); background: var(--warn-bg); }
  .pill.open   { color: var(--ok); }
  .pill.past   { color: var(--ink-faint); }

  /* ---------- money strip ---------- */
  .money { display: grid; gap: 1px; background: var(--line-soft);
           grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); }
  .cell { background: var(--surface); padding: 11px 14px; }
  .cell .k { font-size: 10px; letter-spacing: .08em; text-transform: uppercase;
             color: var(--ink-faint); font-weight: 600; }
  .cell .v { font-family: 'IBM Plex Mono', monospace; font-size: 14.5px;
             font-variant-numeric: tabular-nums; margin-top: 4px;
             word-break: break-word; }
  .cell.flagged { background: var(--warn-bg); }
  .cell.missing .v { color: var(--ink-faint); font-style: italic;
                     font-family: 'IBM Plex Sans', sans-serif; font-size: 13px; }
  .mark { font-size: 11px; font-weight: 700; color: var(--warn);
          cursor: help; margin-left: 4px; }

  /* ---------- meta + detail ---------- */
  .meta { display: grid; gap: 10px 22px; padding: 13px 18px;
          border-top: 1px solid var(--line-soft);
          grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); }
  .meta .k { font-size: 10px; letter-spacing: .08em; text-transform: uppercase;
             color: var(--ink-faint); font-weight: 600; }
  .meta .v { font-size: 13.5px; margin-top: 2px; }
  details { border-top: 1px solid var(--line-soft); }
  summary { cursor: pointer; padding: 10px 18px; font-size: 13px; font-weight: 600;
            color: var(--accent); list-style: none; display: flex; gap: 8px;
            align-items: center; }
  summary::-webkit-details-marker { display: none; }
  summary::before { content: '+'; font-family: 'IBM Plex Mono', monospace;
                    font-weight: 600; width: 12px; }
  details[open] summary::before { content: '\\2212'; }
  .body { padding: 2px 18px 16px; font-size: 13.5px; color: var(--ink-soft);
          white-space: pre-wrap; max-height: 340px; overflow-y: auto;
          border-left: 2px solid var(--accent-soft); margin: 0 18px 14px; }
  .src { font-family: 'IBM Plex Mono', monospace; font-size: 11px;
         color: var(--ink-faint); padding: 0 18px 12px; }

  .empty { padding: 40px; text-align: center; color: var(--ink-faint);
           background: var(--surface); border: 1px dashed var(--line);
           border-radius: 4px; }
  .note { margin: 20px 0 0; padding: 13px 16px; border-radius: 4px;
          background: var(--surface); border: 1px solid var(--line-soft);
          font-size: 13px; color: var(--ink-soft); }
  .note b { color: var(--ink); }
  .legend { display: flex; flex-wrap: wrap; gap: 16px; margin-top: 10px;
            font-size: 12.5px; color: var(--ink-soft); }
  .sw { display: inline-block; width: 11px; height: 11px; border-radius: 2px;
        margin-right: 6px; vertical-align: -1px; }
  footer { margin-top: 34px; padding-top: 14px; border-top: 1px solid var(--line);
           font-size: 12px; color: var(--ink-faint); }
  @media (prefers-reduced-motion: reduce) {
    * { transition: none !important; animation: none !important; }
  }
  @media print {
    body { background: #fff; } .controls, footer { display: none; }
    .card { break-inside: avoid; box-shadow: none; }
    details { display: block; } .body { max-height: none; }
  }
</style>
"""

DASH_BODY = """
<div class="wrap">
  <header class="mast">
    <div class="brand">
      <svg class="mark" width="42" height="42" viewBox="0 0 42 42" aria-hidden="true">
        <rect width="42" height="42" rx="9" fill="var(--accent)"/>
        <text x="21" y="24" text-anchor="middle" font-family="'Source Serif 4', Georgia, serif"
              font-size="17" font-weight="700" fill="var(--surface)">TP</text>
      </svg>
      <div>
        <h1>Tender Pipeline</h1>
        <div class="sub">__SUBTITLE__</div>
      </div>
    </div>
    <div class="runstamp">__STAMP__</div>
  </header>

  <section class="tiles" id="tiles"></section>

  <div class="controls">
    <input type="search" id="q" placeholder="Search tender, place, scope, audit type&hellip;"
           aria-label="Search tenders">
    <label class="chk"><input type="checkbox" id="fRev"> Needs review</label>
    <label class="chk"><input type="checkbox" id="fSoon"> Closing in 14 days</label>
    <select id="sort" aria-label="Sort tenders">
      <option value="due">Sort: deadline first</option>
      <option value="folder">Sort: folder number</option>
      <option value="emd">Sort: EMD highest</option>
      <option value="review">Sort: most to review</option>
    </select>
    <span class="count" id="count"></span>
  </div>

  <main class="cards" id="cards"></main>

  <div class="note">
    <b>How to read this.</b> Every figure carries the file and page it came from &mdash;
    open <em>Where each value came from</em> on any card to check it against the
    document. Amber means the two readers disagreed, confidence was low, or the
    value came off a scanned page. This removes the typing, not the review:
    confirm EMD, fees and the deadline before you act on them.
    <div class="legend">
      <span><span class="sw" style="background:var(--crit)"></span>Closes within 3 days</span>
      <span><span class="sw" style="background:var(--warn)"></span>Closes within 14 days</span>
      <span><span class="sw" style="background:var(--ok)"></span>Open</span>
      <span><span class="sw" style="background:var(--ink-faint)"></span>Closed or no date found</span>
    </div>
  </div>

  <footer>__FOOTER__</footer>
</div>

<script>
const DATA = __DATA__;
const MONEY = ["emd","sd","tender_fees","estimated_cost","assignment_fees"];
const MONEY_LABEL = {emd:"Tender EMD", sd:"Tender SD", tender_fees:"Tender fees",
  estimated_cost:"Estimated cost", assignment_fees:"Assignment fees"};
const PROSE = [["scope_of_work","Scope of work"],["eligibility","Eligibility criteria"],
  ["penalty","Penalty"]];
const MISSING = /^NOT FOUND/i;

const esc = s => String(s==null?"":s).replace(/[&<>"']/g,
  c => ({"&":"&amp;","<":"&lt;",">":"&gt;",'"':"&quot;","'":"&#39;"}[c]));
const has = f => f && f.value && !MISSING.test(f.value);

function parseDue(s){
  if(!s) return null;
  let m = s.match(/(\\d{1,2})[-\\/.](\\d{1,2})[-\\/.](\\d{2,4})/);
  if(m){ let y=+m[3]; if(y<100) y+=2000; return new Date(y, +m[2]-1, +m[1]); }
  m = s.match(/(\\d{1,2})\\s+([A-Za-z]{3,})\\s+(\\d{4})/);
  if(m){ const d=new Date(m[2]+" "+m[1]+", "+m[3]); return isNaN(d)?null:d; }
  return null;
}
function toNum(s){
  if(!s) return 0;
  const m = String(s).replace(/,/g,"").match(/(\\d+(?:\\.\\d+)?)/);
  if(!m) return 0;
  let v = parseFloat(m[1]);
  const tail = String(s).slice(String(s).indexOf(m[1])+m[1].length, 40).toLowerCase();
  if(/^\\s*(?:\\/-)?\\s*crore/.test(tail)) v*=1e7;
  else if(/^\\s*(?:\\/-)?\\s*(?:lakh|lac)/.test(tail)) v*=1e5;
  return v;
}
const DAY = 864e5;
const today = new Date(); today.setHours(0,0,0,0);
function daysLeft(t){
  const d = parseDue(t.fields.submission_date && t.fields.submission_date.value);
  return d ? Math.round((d - today)/DAY) : null;
}
function state(n){
  if(n===null) return "past";
  if(n < 0) return "past";
  if(n <= 3) return "urgent";
  if(n <= 14) return "soon";
  return "open";
}
function reviewCount(t){
  return Object.values(t.fields).filter(f => f.flag).length;
}
const fmtINR = v => v >= 1e7 ? "Rs " + (v/1e7).toFixed(2) + " cr"
                  : v >= 1e5 ? "Rs " + (v/1e5).toFixed(2) + " lakh"
                  : "Rs " + v.toLocaleString("en-IN");

DATA.forEach(t => {
  t._days = daysLeft(t); t._state = state(t._days);
  t._rev = reviewCount(t); t._emd = toNum(has(t.fields.emd) ? t.fields.emd.value : "");
  t._hay = Object.values(t.fields).map(f => f.value).join(" ").toLowerCase()
           + " " + t.folder.toLowerCase();
});

const ICONS = {
  stack: '<svg viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round" stroke-linejoin="round"><rect x="4" y="3" width="12" height="14" rx="1.5"/><line x1="7" y1="7" x2="13" y2="7"/><line x1="7" y1="10.5" x2="13" y2="10.5"/><line x1="7" y1="14" x2="11" y2="14"/></svg>',
  calendar: '<svg viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="4.5" width="14" height="12" rx="1.5"/><line x1="3" y1="8" x2="17" y2="8"/><line x1="6.5" y1="3" x2="6.5" y2="6"/><line x1="13.5" y1="3" x2="13.5" y2="6"/><circle cx="10" cy="12.2" r="1.3" fill="currentColor" stroke="none"/></svg>',
  rupee: '<svg viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.6"><circle cx="10" cy="10" r="7.2"/><text x="10" y="14" text-anchor="middle" font-size="9.5" font-weight="700" fill="currentColor" stroke="none" font-family="inherit">\\u20b9</text></svg>',
  flag: '<svg viewBox="0 0 20 20" fill="none" stroke="currentColor" stroke-width="1.6" stroke-linecap="round" stroke-linejoin="round"><polygon points="10,3 17.5,15.5 2.5,15.5"/><line x1="10" y1="8" x2="10" y2="11.3"/><circle cx="10" cy="13.4" r=".9" fill="currentColor" stroke="none"/></svg>',
};

function tiles(list){
  const live = list.filter(t => t._days !== null && t._days >= 0);
  const next = live.slice().sort((a,b) => a._days - b._days)[0];
  const emd = list.reduce((s,t) => s + t._emd, 0);
  const rev = list.reduce((s,t) => s + t._rev, 0);
  const nodoc = list.filter(t => Object.values(t.fields).every(f => !has(f))).length;
  const t = [
    ["stack", "Tenders", list.length, list.length === 1 ? "in this run" : "folders read"],
    ["calendar", "Next deadline", next ? next._days + (next._days === 1 ? " day" : " days")
      : "\\u2014", next ? next.folder : "none still open", next && next._days <= 3],
    ["rupee", "EMD at stake", emd ? fmtINR(emd) : "\\u2014", "total across open tenders"],
    ["flag", "Fields to verify", rev, "amber or red cells", rev > 0],
  ];
  if (nodoc) t.push(["flag", "No documents", nodoc, "folder(s) with nothing readable", true]);
  document.getElementById("tiles").innerHTML = t.map(([icon,k,v,n,a]) =>
    `<div class="tile${a ? " alert" : ""}"><div class="icon">${ICONS[icon]}</div>
     <div><div class="k">${esc(k)}</div>
     <div class="v">${esc(v)}</div><div class="n">${esc(n)}</div></div></div>`).join("");
}

function cell(t, key){
  const f = t.fields[key] || {};
  const ok = has(f);
  const cls = !ok ? "cell missing" : (f.flag ? "cell flagged" : "cell");
  const mark = f.flag ? `<span class="mark" title="${esc(f.flag)}">&#9888;</span>` : "";
  return `<div class="${cls}"><div class="k">${esc(MONEY_LABEL[key])}</div>
    <div class="v">${ok ? esc(f.value) : "not stated"}${mark}</div></div>`;
}

function card(t){
  const f = t.fields;
  const dueTxt = has(f.submission_date) ? f.submission_date.value : "no date found";
  const pill = t._days === null ? "no deadline read"
    : t._days < 0 ? "closed" : t._days === 0 ? "closes today"
    : t._days + (t._days === 1 ? " day left" : " days left");
  const chips = [];
  if (has(f.purpose)) chips.push(`<span class="chip">${esc(f.purpose.value.slice(0,90))}</span>`);
  if (has(f.period)) chips.push(`<span class="chip plain">Period: ${esc(f.period.value.slice(0,80))}</span>`);
  if (t._rev) chips.push(`<span class="chip plain">${t._rev} to verify</span>`);

  const prose = PROSE.filter(([k]) => has(f[k])).map(([k,lab]) =>
    `<details><summary>${esc(lab)}</summary>
      <div class="body">${esc(f[k].value)}</div>
      <div class="src">Source: ${esc(f[k].ref || "\\u2014")}</div></details>`).join("");

  const refs = Object.entries(f).map(([k,v]) =>
      has(v) ? `${esc(v.label)} &rarr; ${esc(v.ref || "no page recorded")}${
        v.flag ? ` <b>(${esc(v.flag)})</b>` : ""}` : null)
    .filter(Boolean).join("<br>");
  const docs = (t.files || []).map(d =>
    `${esc(d.name)} &mdash; ${d.pages} page(s)${d.ocr ? `, ${d.ocr} via OCR` : ""}${
      d.note ? ` <i>(${esc(d.note)})</i>` : ""}`).join("<br>");

  return `<article class="card ${t._state}">
    <div class="chead">
      <div class="grow">
        <div class="id">${esc(t.folder)}</div>
        <h2>${esc(has(f.tender_name) ? f.tender_name.value : t.folder)}</h2>
        <div class="chips">${chips.join("")}</div>
      </div>
      <div class="due">
        <div class="d">${esc(dueTxt)}</div>
        <span class="pill ${t._state}">${esc(pill)}</span>
      </div>
    </div>
    <div class="money">${MONEY.map(k => cell(t,k)).join("")}</div>
    <div class="meta">
      <div><div class="k">Location / address</div>
        <div class="v">${has(f.location) ? esc(f.location.value) : "<i>not stated</i>"}</div></div>
      <div><div class="k">Period</div>
        <div class="v">${has(f.period) ? esc(f.period.value) : "<i>not stated</i>"}</div></div>
    </div>
    ${prose}
    <details><summary>Where each value came from</summary>
      <div class="body">${refs || "nothing extracted"}</div></details>
    <details><summary>Documents read (${(t.files || []).length})</summary>
      <div class="body">${docs || "none"}</div></details>
  </article>`;
}

function render(){
  const q = document.getElementById("q").value.trim().toLowerCase();
  const rev = document.getElementById("fRev").checked;
  const soon = document.getElementById("fSoon").checked;
  const sort = document.getElementById("sort").value;
  let list = DATA.filter(t => (!q || t._hay.includes(q))
    && (!rev || t._rev > 0)
    && (!soon || (t._days !== null && t._days >= 0 && t._days <= 14)));
  const rank = {urgent:0, soon:1, open:2, past:3};
  list.sort((a,b) => sort === "folder" ? a.folder.localeCompare(b.folder)
    : sort === "emd" ? b._emd - a._emd
    : sort === "review" ? b._rev - a._rev
    : (rank[a._state] - rank[b._state]) || ((a._days ?? 1e9) - (b._days ?? 1e9)));
  document.getElementById("cards").innerHTML = list.length
    ? list.map(card).join("")
    : `<div class="empty">No tender matches that filter.</div>`;
  document.getElementById("count").textContent =
    list.length + " of " + DATA.length + " shown";
  tiles(DATA);
}
["q","fRev","fSoon","sort"].forEach(id =>
  document.getElementById(id).addEventListener("input", render));
render();

// Lets a host page embedding this as an iframe (e.g. the Colab notebook)
// resize the frame to fit, instead of showing a nested scrollbar. Harmless
// when opened as a plain file - there is simply no parent listening.
(function () {
  function report() {
    try { parent.postMessage({ tenderDashHeight: document.body.scrollHeight }, "*"); }
    catch (e) {}
  }
  window.addEventListener("resize", report);
  new MutationObserver(report).observe(document.body, {
    childList: true, subtree: true, attributes: true });
  report();
  setTimeout(report, 300);
})();
</script>
"""


def build_dashboard_html(rows: list, standalone: bool = True,
                         subtitle: str = "", stamp: str = "",
                         footer: str = "") -> str:
    """Render the tender rows as a self-contained HTML dashboard."""
    data = []
    for row in rows:
        res = row["results"]
        data.append({
            "folder": row["tender"],
            "files": row.get("files", []),
            "fields": {
                key: {"label": label, "value": r.value, "ref": r.ref,
                      "conf": r.conf, "flag": r.flag}
                for (key, label), r in ((kl, res[kl[0]]) for kl in FIELDS)
            },
        })
    body = (DASH_BODY
            .replace("__DATA__", json.dumps(data, ensure_ascii=False))
            .replace("__SUBTITLE__", subtitle or
                     f"{len(rows)} tender folders &middot; 13 fields each "
                     f"&middot; every figure traceable to a page")
            .replace("__STAMP__", stamp)
            .replace("__FOOTER__", footer or
                     "Generated by the tender extractor. Values are read from the "
                     "tender documents themselves; always verify before bidding."))
    head = "<title>Tender Pipeline</title>" + DASH_STYLE
    if not standalone:
        return head + body
    return ('<!doctype html>\n<html lang="en">\n<head>\n'
            '<meta charset="utf-8">\n'
            '<meta name="viewport" content="width=device-width, initial-scale=1">\n'
            + head + "\n</head>\n<body>\n" + body + "\n</body>\n</html>\n")


def write_dashboard(rows: list, out_path: Path, stamp: str = "") -> Path:
    flagged = sum(1 for r in rows for k, _ in FIELDS if r["results"][k].flag)
    footer = (f"{len(rows)} tenders &middot; {flagged} field(s) flagged for review "
              f"&middot; open the Excel for the full table")
    html = build_dashboard_html(rows, standalone=True, stamp=stamp, footer=footer)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(html, encoding="utf-8")
    return out_path


if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(description="Extract tender summaries.")
    ap.add_argument("--drive", help="Google Drive folder link (shared, viewer)")
    ap.add_argument("--folder", help="Local folder of tender sub-folders")
    ap.add_argument("--out", default="Tender_Summary.xlsx")
    ap.add_argument("--html", default="Tender_Dashboard.html")
    ap.add_argument("--key", default="", help="Gemini API key (optional)")
    ap.add_argument("--no-ai", action="store_true", help="rules only")
    ap.add_argument("--no-cache", action="store_true")
    a = ap.parse_args()

    if a.folder:
        CONFIG["source_mode"] = "local_folder"
        CONFIG["local_folder"] = a.folder
    elif a.drive:
        CONFIG["source_mode"] = "drive_link"
        CONFIG["drive_folder_url"] = a.drive
    CONFIG["output_xlsx"] = a.out
    CONFIG["output_html"] = a.html
    CONFIG["gemini_api_key"] = a.key or os.environ.get("GEMINI_API_KEY", "")
    CONFIG["use_gemini"] = not a.no_ai
    CONFIG["use_cache"] = not a.no_cache
    run()
